# ==========================================
# EXPEDIENTES/VIEWS.PY - VERSIÓN AMAZON S3 (AWS)
# ==========================================
import io
import os
import json
import uuid
import zipfile
import re
import base64
import logging 
from functools import wraps 
import urllib.parse
# Librerías de seguridad
import bleach
import magic
from django.http import HttpResponse

from .models import Plantilla
from django.utils.text import slugify
import locale
from .models import FacturaGasto
from .utils import procesar_xml_factura
from django.db.models.functions import TruncMonth
from io import BytesIO
from datetime import datetime, timedelta
from decimal import Decimal
from django.utils.formats import date_format

# --- Django Core ---
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt 
from django.contrib import messages
from django.db import transaction
from django.db.models import Count, Q, Sum, Prefetch
from django.http import JsonResponse, HttpResponse, FileResponse
from django.core.files.base import ContentFile
from django.utils import timezone
from django.utils.html import strip_tags
from django.template.loader import render_to_string
from django.core.mail import EmailMultiAlternatives, EmailMessage, send_mail
from django.conf import settings

# --- Librerías de Terceros ---
import pandas as pd
import mammoth
from docxtpl import DocxTemplate
from docx import Document as DocumentoWord
import weasyprint
import qrcode
from email.mime.image import MIMEImage

# --- Modelos Locales ---
from .models import (
    Usuario, Cliente, Carpeta, Expediente, Documento, 
    Tarea, Bitacora, Plantilla, VariableEstandar,
    Servicio, Cotizacion, ItemCotizacion, PlantillaMensaje,
    CuentaPorCobrar, Pago, Evento, CampoAdicional, Archivo,
    SolicitudEnlace, ArchivoTemporal
)

# --- Utilidades ---
from .utils import generar_pdf_response

logger = logging.getLogger(__name__)

EMAIL_REPLY_TO = os.environ.get('EMAIL_REPLY_TO', 'maribel.aldana@gestionescorpad.com')
URL_PORTAL = os.environ.get('URL_PORTAL', 'https://portalgestionescorpad.up.railway.app')
FIRMA_NOMBRE_DEFAULT = os.environ.get('FIRMA_NOMBRE_DEFAULT', 'Lic. Maribel Aldana Santos')
FIRMA_CARGO_DEFAULT = os.environ.get('FIRMA_CARGO_DEFAULT', 'Gestiones Corpad | Directora General')


# <--- DECORADOR CENTRALIZADO DE PERMISOS ---
def requiere_permiso(permiso):
    def decorator(view_func):
        @wraps(view_func)
        def _wrapped_view(request, *args, **kwargs):
            if request.user.rol == 'admin' or getattr(request.user, permiso, False):
                return view_func(request, *args, **kwargs)
            messages.error(request, f"Acceso Denegado. Se requiere el permiso: {permiso}")
            return redirect('dashboard')
        return _wrapped_view
    return decorator


# <--- HELPER CENTRALIZADO DE BITÁCORA ---
def registrar_bitacora(usuario, cliente, accion, descripcion):
    try:
        Bitacora.objects.create(
            usuario=usuario,
            cliente=cliente,
            accion=accion,
            descripcion=descripcion
        )
    except Exception as e:
        logger.warning(f"No se pudo registrar bitácora [{accion}]: {e}")


# ==========================================
# 1. AUTENTICACIÓN Y PERFIL
# ==========================================

def signout(request):
    logout(request)
    return redirect('login')

def registro(request):
    if request.method == 'POST':
        data = request.POST
        if data.get('pass1') != data.get('pass2'):
            messages.error(request, "Las contraseñas no coinciden.")
            return render(request, 'registro.html')

        if Usuario.objects.filter(username=data.get('username')).exists():
            messages.error(request, "El usuario ya existe.")
            return render(request, 'registro.html')

        try:
            Usuario.objects.create_user(
                username=data.get('username'), 
                email=data.get('email'), 
                password=data.get('pass1'),
                first_name=data.get('first_name'), 
                last_name=data.get('last_name'), 
                is_active=False
            )
            return render(request, 'registro_pendiente.html')
        except Exception as e:
            logger.error(f"Error en registro de usuario: {e}")
            messages.error(request, f"Error del sistema: {e}")

    return render(request, 'registro.html')

@login_required
def mi_perfil(request):
    user = request.user
    if request.method == 'POST':
        user.first_name = request.POST.get('first_name')
        user.last_name = request.POST.get('last_name')
        user.email = request.POST.get('email')
        user.telefono = request.POST.get('telefono')
        user.puesto = request.POST.get('puesto')
        
        if request.FILES.get('avatar'):
            user.avatar = request.FILES['avatar']
            
        user.save()
        messages.success(request, "Perfil actualizado correctamente.")
        return redirect('mi_perfil')
    return render(request, 'usuarios/mi_perfil.html', {'user': user})

# ==========================================
# 2. GESTIÓN DE USUARIOS (ADMIN)
# ==========================================

@login_required
def gestion_usuarios(request):
    if request.user.rol != 'admin': 
        return redirect('dashboard')
    usuarios = Usuario.objects.all().order_by('-date_joined')
    return render(request, 'gestion_usuarios.html', {'usuarios': usuarios})

@login_required
def autorizar_usuario(request, user_id):
    if request.user.rol != 'admin': 
        return redirect('dashboard')
    user = get_object_or_404(Usuario, id=user_id)
    user.is_active = True
    user.save()
    messages.success(request, f"Usuario {user.username} autorizado.")
    return redirect('gestion_usuarios')

@login_required
def editar_usuario(request, user_id):
    if request.user.rol != 'admin': 
        return redirect('dashboard')
    user_obj = get_object_or_404(Usuario, id=user_id)
    clientes_disponibles = Cliente.objects.all().order_by('nombre_empresa')
    
    if request.method == 'POST':
        user_obj.rol = request.POST.get('rol')
        user_obj.first_name = request.POST.get('first_name') or ""
        user_obj.last_name = request.POST.get('last_name') or ""
        user_obj.email = request.POST.get('email')
        user_obj.telefono = request.POST.get('telefono') or None
        user_obj.puesto = request.POST.get('puesto') or None
        
        permisos = ['can_create_client', 'can_edit_client', 'can_delete_client', 
                    'can_upload_files', 'can_view_documents', 'can_manage_users',
                    'access_finanzas', 'access_cotizaciones', 'access_contratos', 
                    'access_disenador', 'access_agenda', 'access_gastos', 'access_qr']
        
        for p in permisos:
            setattr(user_obj, p, request.POST.get(p) == 'on')
        
        clientes_ids = request.POST.getlist('clientes_asignados')
        user_obj.save()
        
        if user_obj.rol != 'admin':
            user_obj.clientes_asignados.set(clientes_ids)
        else:
            user_obj.clientes_asignados.clear()
            
        messages.success(request, f"Permisos de {user_obj.username} actualizados.")
        return redirect('gestion_usuarios')

    return render(request, 'usuarios/editar_usuario.html', {'u': user_obj, 'clientes': clientes_disponibles})

@login_required
def eliminar_usuario(request, user_id):
    if request.user.rol != 'admin': 
        return redirect('dashboard')
    u = get_object_or_404(Usuario, id=user_id)
    if u == request.user:
        messages.error(request, "No puedes eliminarte a ti mismo.")
        return redirect('gestion_usuarios')
    u.delete()
    messages.success(request, "Usuario eliminado.")
    return redirect('gestion_usuarios')

# ==========================================
# 3. DASHBOARD Y CLIENTES
# ==========================================

@login_required
def dashboard(request):
    qs = Cliente.objects.annotate(
        num_expedientes=Count('expedientes', distinct=True),
        urgencias=Count('tareas', filter=Q(tareas__prioridad='alta', tareas__completada=False), distinct=True)
    ).order_by('-urgencias', '-fecha_registro')

    if request.user.rol == 'admin':
        mis_clientes = qs
        base_clientes = Cliente.objects.all()
    else:
        mis_clientes = qs.filter(abogados_asignados=request.user)
        base_clientes = request.user.clientes_asignados.all()

    stats = {
        'total_clientes': base_clientes.count(),
        'expedientes_activos': Expediente.objects.filter(cliente__in=base_clientes, estado='abierto').count(),
        'tareas_pendientes': Tarea.objects.filter(cliente__in=base_clientes, completada=False).count(),
        'docs_subidos': Documento.objects.filter(cliente__in=base_clientes).count()
    }
    
    hoy = timezone.now().date()
    tareas_criticas = Tarea.objects.filter(
        cliente__in=base_clientes, 
        completada=False, 
        fecha_limite__lte=hoy
    ).select_related('cliente')
    
    pendientes = 0
    if request.user.rol == 'admin':
        pendientes = Usuario.objects.filter(is_active=False).count()

    return render(request, 'dashboard.html', {
        'clientes': mis_clientes,
        'stats': stats,
        'usuarios_pendientes_conteo': pendientes,
        'now': timezone.now(),
        'alertas': {'tareas': tareas_criticas} 
    })


@login_required
@requiere_permiso('can_create_client')
def nuevo_cliente(request):
    if request.method == 'POST':
        c = Cliente.objects.create(
            nombre_empresa=request.POST.get('nombre_empresa'),
            nombre_contacto=request.POST.get('nombre_contacto'),
            email=request.POST.get('email'),
            telefono=request.POST.get('telefono'),
            logo=request.FILES.get('logo')
        )
        if request.user.rol != 'admin':
            request.user.clientes_asignados.add(c)

        # Lógica para conservar solo las carpetas seleccionadas
        carpetas_seleccionadas = request.POST.getlist('carpetas_seleccionadas')
        carpetas_base = [
            'CARPETA ADMINISTRATIVA', 'LICENCIA DE FUNCIONAMIENTO', 
            'PROGRAMA ESPECIFICO DE PROTECCIÓN CIVIL', 'PROTECCIÓN CIVIL MUNICIPAL', 
            'PROTECCIÓN CIVIL ESTATAL', 'MEDIO AMBIENTE', 'REGISTRO AMBIENTAL ESTATAL', 
            'CEDULA DE ZONIFICACIÓN', 'LICENCIA DE USO DE SUELO'
        ]
        
        for nombre_carpeta in carpetas_base:
            if nombre_carpeta not in carpetas_seleccionadas:
                Carpeta.objects.filter(cliente=c, nombre=nombre_carpeta).delete()

        registrar_bitacora(request.user, c, 'creacion', f"Dio de alta al cliente '{c.nombre_empresa}'.")

        # --- MAGIA: SINCRONIZACIÓN AUTOMÁTICA DE SUCURSALES (VERSIÓN AVANZADA) ---
        if c.logo:
            palabras = c.nombre_empresa.upper().split()
            
            if len(palabras) > 0:
                palabras_genericas = [
                    'GRUPO', 'OPERADORA', 'COMERCIALIZADORA', 'EL', 'LA', 'LOS', 'LAS', 
                    'CORPORATIVO', 'CONSORCIO', 'GASTRONOMIA', 'SERVICIOS', 'CONSTRUCTORA', 
                    'PROMOTORA', 'PROVEEDORA', 'DISTRIBUIDORA', 'INMOBILIARIA', 'TRANSPORTES', 
                    'LOGISTICA', 'RESTAURANTE', 'HOTEL', 'CLINICA', 'HOSPITAL', 'INSTITUTO', 
                    'COLEGIO', 'AGENCIA', 'DESPACHO', 'ASOCIACION', 'SOCIEDAD', 'SISTEMAS', 
                    'INDUSTRIAS', 'ADMINISTRADORA', 'CENTRO', 'FABRICA', 'PRODUCTORA'
                ]
                
                # Si empieza con palabra genérica, buscamos coincidencias con las primeras dos palabras
                if palabras[0] in palabras_genericas and len(palabras) > 1:
                    clave_busqueda = f"{palabras[0]} {palabras[1]}"
                else:
                    # Si es marca propia, buscamos coincidencias solo con la primera palabra
                    clave_busqueda = palabras[0]
                
                if len(clave_busqueda) > 3:
                    sucursales = Cliente.objects.filter(nombre_empresa__istartswith=clave_busqueda).exclude(id=c.id)
                    
                    count_actualizadas = 0
                    for sucursal in sucursales:
                        sucursal.logo = c.logo
                        sucursal.save(update_fields=['logo'])
                        count_actualizadas += 1
                    
                    if count_actualizadas > 0:
                        messages.info(request, f"💡 Inteligencia AppLegal: El logo se aplicó automáticamente a {count_actualizadas} sucursal(es) de '{clave_busqueda}'.")

        return redirect('dashboard')
    return render(request, 'nuevo_cliente.html')

@login_required
@requiere_permiso('can_delete_client')
def eliminar_cliente(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    nombre = cliente.nombre_empresa
    registrar_bitacora(request.user, cliente, 'eliminacion', f"Eliminó el cliente '{nombre}' y todos sus datos.")
    cliente.delete()
    messages.success(request, "Cliente eliminado.")
    return redirect('dashboard')

@login_required
def detalle_cliente(request, cliente_id, carpeta_id=None):
    # 1. Seguridad: Verificar si el usuario tiene acceso a este cliente
    if request.user.rol != 'admin':
        cliente = get_object_or_404(request.user.clientes_asignados, id=cliente_id)
    else:
        cliente = get_object_or_404(Cliente, id=cliente_id)

    # 2. Navegación (Breadcrumbs)
    carpeta_actual = None
    breadcrumbs = []
    if carpeta_id:
        carpeta_actual = get_object_or_404(Carpeta, id=carpeta_id, cliente=cliente)
        
        # Construir breadcrumbs (Ruta de navegación)
        temp = carpeta_actual.padre
        while temp:
            breadcrumbs.insert(0, temp)
            temp = temp.padre

    # 3. Obtener subcarpetas
    if carpeta_actual:
        carpetas = Carpeta.objects.filter(cliente=cliente, padre=carpeta_actual).order_by('nombre')
    else:
        carpetas = Carpeta.objects.filter(cliente=cliente, padre__isnull=True).order_by('nombre')

    # 4. Obtener archivos
    if carpeta_actual:
        documentos = Documento.objects.filter(cliente=cliente, carpeta=carpeta_actual).order_by('-fecha_subida')
    else:
        documentos = Documento.objects.filter(cliente=cliente, carpeta__isnull=True).order_by('-fecha_subida')

    # 5. Obtener Historial de Bitácora (Solo últimos 50 para no saturar)
    historial = Bitacora.objects.filter(cliente=cliente).select_related('usuario').order_by('-fecha')[:50]

    # 6. Estadísticas Generales del Cliente
    total_docs = Documento.objects.filter(cliente=cliente).count()
    carpetas_base = Carpeta.objects.filter(cliente=cliente, padre__isnull=True).count()

    # 7. Obtener todas las carpetas (para el Modal de Mover Archivo)
    todas_carpetas = Carpeta.objects.filter(cliente=cliente).order_by('nombre')
    
    # 8. Obtener archivos temporales pendientes de auditoría
    archivos_pendientes = ArchivoTemporal.objects.filter(
        solicitud__cliente=cliente
    ).order_by('-fecha_subida')

    # ---> 9. INTELIGENCIA: Contar sucursales relacionadas para el Modal de Borrado Masivo <---
    sucursales_relacionadas_count = 1
    clave_busqueda = ""
    if cliente.nombre_empresa:
        palabras = cliente.nombre_empresa.upper().split()
        if len(palabras) > 0:
            palabras_genericas = [
                'GRUPO', 'OPERADORA', 'COMERCIALIZADORA', 'EL', 'LA', 'LOS', 'LAS', 
                'CORPORATIVO', 'CONSORCIO', 'GASTRONOMIA', 'SERVICIOS', 'CONSTRUCTORA', 
                'PROMOTORA', 'PROVEEDORA', 'DISTRIBUIDORA', 'INMOBILIARIA', 'TRANSPORTES', 
                'LOGISTICA', 'RESTAURANTE', 'HOTEL', 'CLINICA', 'HOSPITAL', 'INSTITUTO', 
                'COLEGIO', 'AGENCIA', 'DESPACHO', 'ASOCIACION', 'SOCIEDAD', 'SISTEMAS', 
                'INDUSTRIAS', 'ADMINISTRADORA', 'CENTRO', 'FABRICA', 'PRODUCTORA'
            ]
            if palabras[0] in palabras_genericas and len(palabras) > 1:
                clave_busqueda = f"{palabras[0]} {palabras[1]}"
            else:
                clave_busqueda = palabras[0]
            
            if len(clave_busqueda) > 3:
                sucursales_relacionadas_count = Cliente.objects.filter(nombre_empresa__istartswith=clave_busqueda).count()

    # 10. Construcción del Contexto
    context = {
        'cliente': cliente,
        'carpeta_actual': carpeta_actual,
        'carpetas': carpetas,
        'documentos': documentos,
        'breadcrumbs': breadcrumbs,
        'todas_carpetas': todas_carpetas,
        'historial': historial,
        'archivos_pendientes': archivos_pendientes,
        'stats_cliente': {
            'total_docs': total_docs,
            'expedientes_activos': carpetas_base
        },
        'sucursales_relacionadas_count': sucursales_relacionadas_count,
        'clave_busqueda': clave_busqueda,
    }
    return render(request, 'detalle_cliente.html', context)
@login_required
def editar_cliente(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    if request.user.rol != 'admin' and cliente not in request.user.clientes_asignados.all():
        return redirect('dashboard')

    campos_dinamicos = CampoAdicional.objects.all()

    if request.method == 'POST':
        cliente.nombre_empresa = request.POST.get('nombre_empresa')
        cliente.nombre_contacto = request.POST.get('nombre_contacto')
        cliente.email = request.POST.get('email')
        cliente.telefono = request.POST.get('telefono')
        
        logo_actualizado = False
        if request.FILES.get('logo'):
            cliente.logo = request.FILES['logo']
            logo_actualizado = True

        datos_nuevos = cliente.datos_extra or {}
        for campo in campos_dinamicos:
            valor = request.POST.get(f"custom_{campo.id}")
            if valor:
                datos_nuevos[campo.nombre] = valor
        
        cliente.datos_extra = datos_nuevos
        cliente.save()
        registrar_bitacora(request.user, cliente, 'edicion', "Actualizó datos del cliente.")
        messages.success(request, "Cliente actualizado.")

        # --- MAGIA: SINCRONIZACIÓN AUTOMÁTICA DE SUCURSALES (VERSIÓN AVANZADA) ---
        if logo_actualizado and cliente.logo:
            palabras = cliente.nombre_empresa.upper().split()
            
            if len(palabras) > 0:
                palabras_genericas = [
                    'GRUPO', 'OPERADORA', 'COMERCIALIZADORA', 'EL', 'LA', 'LOS', 'LAS', 
                    'CORPORATIVO', 'CONSORCIO', 'GASTRONOMIA', 'SERVICIOS', 'CONSTRUCTORA', 
                    'PROMOTORA', 'PROVEEDORA', 'DISTRIBUIDORA', 'INMOBILIARIA', 'TRANSPORTES', 
                    'LOGISTICA', 'RESTAURANTE', 'HOTEL', 'CLINICA', 'HOSPITAL', 'INSTITUTO', 
                    'COLEGIO', 'AGENCIA', 'DESPACHO', 'ASOCIACION', 'SOCIEDAD', 'SISTEMAS', 
                    'INDUSTRIAS', 'ADMINISTRADORA', 'CENTRO', 'FABRICA', 'PRODUCTORA'
                ]
                
                # Si empieza con palabra genérica, buscamos coincidencias con las primeras dos palabras
                if palabras[0] in palabras_genericas and len(palabras) > 1:
                    clave_busqueda = f"{palabras[0]} {palabras[1]}"
                else:
                    # Si es marca propia, buscamos coincidencias solo con la primera palabra
                    clave_busqueda = palabras[0]
                
                if len(clave_busqueda) > 3:
                    sucursales = Cliente.objects.filter(nombre_empresa__istartswith=clave_busqueda).exclude(id=cliente.id)
                    
                    count_actualizadas = 0
                    for sucursal in sucursales:
                        sucursal.logo = cliente.logo
                        sucursal.save(update_fields=['logo'])
                        count_actualizadas += 1
                    
                    if count_actualizadas > 0:
                        messages.info(request, f"💡 Inteligencia AppLegal: El logo se sincronizó automáticamente con {count_actualizadas} sucursal(es) de '{clave_busqueda}'.")

        return redirect('detalle_cliente', cliente_id=cliente.id)

    return render(request, 'clientes/editar.html', {
        'c': cliente,
        'campos_dinamicos': campos_dinamicos,
        'datos_existentes': cliente.datos_extra
    })
# ==========================================
# 4. CONFIGURACIÓN Y DRIVE
# ==========================================

@login_required
def configurar_campos(request):
    if request.user.rol != 'admin': 
        return redirect('dashboard')
    campos = CampoAdicional.objects.all()
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        if not CampoAdicional.objects.filter(nombre__iexact=nombre).exists():
            CampoAdicional.objects.create(nombre=nombre, tipo=request.POST.get('tipo'))
            messages.success(request, f"Campo '{nombre}' agregado.")
        return redirect('configurar_campos')
    return render(request, 'clientes/configurar_campos.html', {'campos': campos})

@login_required
def eliminar_campo_dinamico(request, campo_id):
    if request.user.rol != 'admin': 
        return redirect('dashboard')
    get_object_or_404(CampoAdicional, id=campo_id).delete()
    return redirect('configurar_campos')

@login_required
def crear_carpeta(request, cliente_id):
    if request.method == 'POST':
        padre_id = request.POST.get('padre_id')
        padre = get_object_or_404(Carpeta, id=padre_id) if padre_id else None
        carpeta = Carpeta.objects.create(nombre=request.POST.get('nombre'), cliente_id=cliente_id, padre=padre)
        cliente = get_object_or_404(Cliente, id=cliente_id)
        ubicacion = f"dentro de '{padre.nombre}'" if padre else "en la raíz del expediente"
        registrar_bitacora(request.user, cliente, 'creacion', f"Creó la carpeta '{carpeta.nombre}' {ubicacion}.")
        if padre: 
            return redirect('detalle_carpeta', cliente_id=cliente_id, carpeta_id=padre.id)
    return redirect('detalle_cliente', cliente_id=cliente_id)

@login_required
@requiere_permiso('can_delete_client')
def eliminar_carpeta(request, carpeta_id):
    c = get_object_or_404(Carpeta, id=carpeta_id)
    cliente = c.cliente
    nombre_carpeta = c.nombre
    url_destino = 'detalle_carpeta' if c.padre else 'detalle_cliente'
    kwargs = {'cliente_id': c.cliente.id}
    if c.padre: 
        kwargs['carpeta_id'] = c.padre.id
    registrar_bitacora(request.user, cliente, 'eliminacion', f"Eliminó la carpeta '{nombre_carpeta}' y todo su contenido.")
    c.delete()
    return redirect(url_destino, **kwargs)

@login_required
def crear_expediente(request, cliente_id):
    if request.method == 'POST':
        cliente = get_object_or_404(Cliente, id=cliente_id)
        num_exp = request.POST.get('num_expediente')
        titulo = request.POST.get('titulo')
        f = Carpeta.objects.create(nombre=f"EXP {num_exp}: {titulo}", cliente_id=cliente_id, es_expediente=True)
        Expediente.objects.create(cliente_id=cliente_id, num_expediente=num_exp, titulo=titulo, carpeta=f)
        registrar_bitacora(request.user, cliente, 'creacion', f"Creó el expediente #{num_exp}: {titulo}.")
    return redirect('detalle_cliente', cliente_id=cliente_id)

@login_required
@requiere_permiso('can_upload_files')
def subir_archivo_drive(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    
    if request.method == 'POST':
        archivos = request.FILES.getlist('archivo')
        carpeta_id = request.POST.get('carpeta_id')
        fecha_vencimiento = request.POST.get('fecha_vencimiento')
        
        carpeta_padre_original = None
        if carpeta_id:
            carpeta_padre_original = get_object_or_404(Carpeta, id=carpeta_id)

        # LÓGICA DE CARPETAS DESDE EL FRONTEND
        rutas_str = request.POST.get('rutas_json', '{}')
        try:
            rutas_dict = json.loads(rutas_str)
        except:
            rutas_dict = {}

        eventos_to_create = []
        archivos_guardados = 0
        nombres_guardados = []

        for f in archivos:
            # ---> SE ELIMINÓ LA RESTRICCIÓN DE MAGIC PARA ADMITIR TODO TIPO DE ARCHIVOS <---

            ruta_relativa = rutas_dict.get(f.name, f.name)
            carpeta_destino = carpeta_padre_original

            # Si el archivo viene dentro de una subcarpeta, la creamos/buscamos
            if '/' in ruta_relativa:
                partes = ruta_relativa.split('/')
                nombres_carpetas = partes[:-1] # Excluye el archivo
                
                for nombre_carpeta in nombres_carpetas:
                    carpeta_destino, created = Carpeta.objects.get_or_create(
                        cliente=cliente,
                        nombre=nombre_carpeta,
                        padre=carpeta_destino
                    )

            nuevo_doc = Documento(
                cliente=cliente,
                carpeta=carpeta_destino,
                archivo=f,
                nombre_archivo=f.name,
                subido_por=request.user
            )
            
            if fecha_vencimiento:
                nuevo_doc.fecha_vencimiento = fecha_vencimiento
                fecha_fin = datetime.strptime(fecha_vencimiento, '%Y-%m-%d').date()
                alertas = [(20, '⚠️ Vence en 20 días'), (10, '🟠 Vence en 10 días'), (5, '🔴 URGENTE: Vence en 5 días')]
                
                for dias_antes, prefijo in alertas:
                    fecha_alerta = fecha_fin - timedelta(days=dias_antes)
                    if fecha_alerta >= timezone.now().date():
                        eventos_to_create.append(Evento(
                            cliente=cliente,
                            usuario=request.user,
                            titulo=f"{prefijo}: {f.name}",
                            inicio=datetime.combine(fecha_alerta, datetime.min.time()),
                            fin=datetime.combine(fecha_alerta, datetime.min.time()) + timedelta(hours=1),
                            descripcion=f"Recordatorio automático de vencimiento para el documento: {f.name}"
                        ))
            
            nuevo_doc.save()
            archivos_guardados += 1
            nombres_guardados.append(f.name)

        if eventos_to_create:
            Evento.objects.bulk_create(eventos_to_create)

        if archivos_guardados > 0:
            ubicacion = f"en '{carpeta_padre_original.nombre}'" if carpeta_padre_original else "en la raíz"
            resumen = ', '.join(nombres_guardados[:5]) + ('...' if len(nombres_guardados) > 5 else '')
            registrar_bitacora(
                request.user, cliente, 'subida',
                f"Subió {archivos_guardados} archivo(s) {ubicacion}: {resumen}"
            )
            messages.success(request, f"{archivos_guardados} archivo(s) subido(s) correctamente.")
        
        if carpeta_padre_original:
            return redirect('detalle_carpeta', cliente_id=cliente.id, carpeta_id=carpeta_padre_original.id)
        return redirect('detalle_cliente', cliente_id=cliente.id)
        
    return redirect('detalle_cliente', cliente_id=cliente.id)

@login_required
@requiere_permiso('can_delete_client')
def eliminar_archivo_drive(request, archivo_id):
    doc = get_object_or_404(Documento, id=archivo_id)
    c_id, padre_id = doc.cliente.id, doc.carpeta.id if doc.carpeta else None
    registrar_bitacora(request.user, doc.cliente, 'eliminacion', f"Eliminó el archivo '{doc.nombre_archivo}'.")
    doc.archivo.delete()
    doc.delete()
    if padre_id: 
        return redirect('detalle_carpeta', cliente_id=c_id, carpeta_id=padre_id)
    return redirect('detalle_cliente', cliente_id=c_id)

@login_required
def descargar_carpeta_zip(request, carpeta_id):
    carpeta = get_object_or_404(Carpeta, id=carpeta_id)
    if request.user.rol != 'admin' and carpeta.cliente not in request.user.clientes_asignados.all():
        return HttpResponse("Acceso Denegado", status=403)
    
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for file in Documento.objects.filter(carpeta=carpeta):
            try: 
                zip_file.writestr(file.nombre_archivo, file.archivo.read())
            except Exception as e: 
                logger.warning(f"No se pudo incluir {file.nombre_archivo} en ZIP de carpeta {carpeta.id}: {e}")
    
    registrar_bitacora(request.user, carpeta.cliente, 'descarga', f"Descargó ZIP completo de la carpeta '{carpeta.nombre}'.")
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{carpeta.nombre}.zip"'
    return response

@login_required
def acciones_masivas_drive(request):
    if request.method == 'POST':
        accion = request.POST.get('accion')
        doc_ids = request.POST.getlist('doc_ids')
        docs = Documento.objects.filter(id__in=doc_ids)
        if not docs: 
            return redirect(request.META.get('HTTP_REFERER'))
        
        cliente = docs.first().cliente

        if accion == 'eliminar':
            if not (request.user.can_delete_client or request.user.rol == 'admin'): 
                return redirect(request.META.get('HTTP_REFERER'))
            count = docs.count()
            nombres = list(docs.values_list('nombre_archivo', flat=True))
            for doc in docs: 
                doc.archivo.delete()
                doc.delete()
            resumen = ', '.join(nombres[:5]) + ('...' if len(nombres) > 5 else '')
            registrar_bitacora(request.user, cliente, 'eliminacion', f"Eliminó {count} archivo(s) masivamente: {resumen}.")
            messages.success(request, f"Se eliminaron {count} archivos.")
        
        elif accion == 'descargar':
            buffer = BytesIO()
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for doc in docs:
                    try: 
                        zip_file.writestr(doc.nombre_archivo, doc.archivo.read())
                    except Exception as e:
                        logger.warning(f"Error zipping {doc.id} en acciones masivas: {e}")
            registrar_bitacora(request.user, cliente, 'descarga', f"Descargó selección de {docs.count()} archivo(s) en ZIP.")
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="Seleccion.zip"'
            return response
            
    return redirect(request.META.get('HTTP_REFERER'))

@login_required
def mover_archivo_drive(request, archivo_id):
    doc = get_object_or_404(Documento, id=archivo_id)
    
    if not (request.user.can_edit_client or request.user.can_upload_files or request.user.rol == 'admin'):
        messages.error(request, "No tienes permiso para mover archivos.")
        return redirect('detalle_cliente', cliente_id=doc.cliente.id)

    if request.method == 'POST':
        destino_id = request.POST.get('carpeta_destino')
        fecha_vencimiento = request.POST.get('fecha_vencimiento')
        origen_nombre = doc.carpeta.nombre if doc.carpeta else "Raíz"
        
        if destino_id == 'ROOT':
            doc.carpeta = None
            nombre_destino = "Carpeta Raíz"
        else:
            carpeta_destino = get_object_or_404(Carpeta, id=destino_id)
            doc.carpeta = carpeta_destino
            nombre_destino = carpeta_destino.nombre
            
        if fecha_vencimiento: 
            doc.fecha_vencimiento = fecha_vencimiento
            fecha_fin = datetime.strptime(fecha_vencimiento, '%Y-%m-%d').date()
            alertas = [(20, '⚠️ Vence en 20 días'), (10, '🟠 Vence en 10 días'), (5, '🔴 URGENTE: Vence en 5 días')]
            eventos_to_create = []
            for dias_antes, prefijo in alertas:
                fecha_alerta = fecha_fin - timedelta(days=dias_antes)
                if fecha_alerta >= timezone.now().date():
                    eventos_to_create.append(Evento(
                        cliente=doc.cliente,
                        usuario=request.user,
                        titulo=f"{prefijo}: {doc.nombre_archivo}",
                        inicio=datetime.combine(fecha_alerta, datetime.min.time()),
                        fin=datetime.combine(fecha_alerta, datetime.min.time()) + timedelta(hours=1),
                        descripcion=f"Recordatorio automático de vencimiento para el documento: {doc.nombre_archivo}"
                    ))
            if eventos_to_create:
                Evento.objects.bulk_create(eventos_to_create)

        doc.save()
        registrar_bitacora(
            request.user, doc.cliente, 'movimiento',
            f"Movió '{doc.nombre_archivo}' de '{origen_nombre}' a '{nombre_destino}'."
        )
        messages.success(request, f"Archivo movido a: {nombre_destino}")
        
    return redirect(request.META.get('HTTP_REFERER', 'dashboard'))

# ==========================================
# 5. TAREAS
# ==========================================

@login_required
def gestionar_tarea(request, cliente_id):
    if request.method == 'POST':
        cliente = get_object_or_404(Cliente, id=cliente_id)
        tarea = Tarea.objects.create(
            cliente_id=cliente_id, 
            titulo=request.POST.get('titulo'),
            fecha_limite=request.POST.get('fecha_limite'), 
            prioridad=request.POST.get('prioridad')
        )
        registrar_bitacora(request.user, cliente, 'tarea', f"Creó la tarea '{tarea.titulo}' con vencimiento {tarea.fecha_limite}.")
    return redirect('detalle_cliente', cliente_id=cliente_id)

@login_required
def toggle_tarea(request, tarea_id):
    t = get_object_or_404(Tarea, id=tarea_id)
    t.completada = not t.completada
    t.save()
    estado = "completó" if t.completada else "reabrió"
    registrar_bitacora(request.user, t.cliente, 'tarea', f"{estado.capitalize()} la tarea '{t.titulo}'.")
    return redirect('detalle_cliente', cliente_id=t.cliente.id)

@login_required
def editar_tarea(request, tarea_id):
    t = get_object_or_404(Tarea, id=tarea_id)
    if request.method == 'POST':
        titulo_anterior = t.titulo
        t.titulo = request.POST.get('titulo')
        t.fecha_limite = request.POST.get('fecha_limite')
        t.prioridad = request.POST.get('prioridad')
        t.save()
        registrar_bitacora(request.user, t.cliente, 'tarea', f"Editó la tarea '{titulo_anterior}' → '{t.titulo}'.")
    return redirect('detalle_cliente', cliente_id=t.cliente.id)

@login_required
def eliminar_tarea(request, tarea_id):
    t = get_object_or_404(Tarea, id=tarea_id)
    c_id = t.cliente.id
    registrar_bitacora(request.user, t.cliente, 'eliminacion', f"Eliminó la tarea '{t.titulo}'.")
    t.delete()
    return redirect('detalle_cliente', cliente_id=c_id)

# ==========================================
# 6. CONTRATOS Y DISEÑADOR
# ==========================================

@login_required
@requiere_permiso('access_contratos')
def generador_contratos(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    
    if request.method == 'GET' and 'plantilla_id' not in request.GET:
        return render(request, 'generador/seleccionar.html', {
            'cliente': cliente,
            'plantillas': Plantilla.objects.all().order_by('-fecha_subida'),
            'glosario': VariableEstandar.objects.all().order_by('clave')
        })

    plantilla = get_object_or_404(Plantilla, id=request.GET.get('plantilla_id') or request.POST.get('plantilla_id'))
    
    doc = DocxTemplate(io.BytesIO(plantilla.archivo.read()))
    
    vars_en_doc = doc.get_undeclared_template_variables()
    memoria = cliente.datos_extra if isinstance(cliente.datos_extra, dict) else {}
    formulario = []
    
    hoy = timezone.now()
    meses = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
    fecha_larga = f"{hoy.day} de {meses[hoy.month-1]} de {hoy.year}"
    
    datos_sistema = {
        'cliente_empresa': cliente.nombre_empresa,
        'cliente_contacto': cliente.nombre_contacto,
        'cliente_email': cliente.email,
        'cliente_telefono': cliente.telefono,
        'cliente_direccion': memoria.get('direccion', ''),
        'cliente_cargo': memoria.get('cargo', ''),
        'fecha_corta': hoy.strftime("%d/%m/%Y"),
        'fecha_larga': fecha_larga,
        'anio_actual': str(hoy.year),
        'firma_nombre': FIRMA_NOMBRE_DEFAULT,
        'firma_cargo': FIRMA_CARGO_DEFAULT,
    }

    cotizacion = Cotizacion.objects.filter(cliente_convertido=cliente, estado='aceptada').last()
    if cotizacion:
        datos_sistema.update({
            'monto_total': f"${cotizacion.total_con_iva:,.2f}",
            'monto_subtotal': f"${cotizacion.total:,.2f}",
            'monto_anticipo': f"${(cotizacion.total_con_iva/2):,.2f}" if cotizacion.condiciones_pago == '50_50' else 'N/A',
            'proyecto_titulo': cotizacion.titulo or 'Gestión Administrativa'
        })

    vars_std_dict = {v.clave: v for v in VariableEstandar.objects.all()}

    for v in vars_en_doc:
        if v in datos_sistema:
            formulario.append({'clave': v, 'valor': datos_sistema[v], 'descripcion': 'Automático del Sistema', 'es_automatico': True, 'tipo': 'hidden'})
            continue

        var_std = vars_std_dict.get(v)
        val = ""
        desc = "Variable personalizada"
        tipo = "text"

        if var_std:
            desc = var_std.descripcion
            if var_std.tipo == 'fecha': tipo = 'date'
            val = memoria.get(v, '')
        else: 
            val = memoria.get(v, '')

        formulario.append({'clave': v, 'valor': val, 'descripcion': desc, 'es_automatico': False, 'tipo': tipo})

    if request.method == 'POST':
        contexto = {}
        nuevos_datos = {}
        for item in formulario:
            if item['es_automatico']: 
                val = item['valor'] 
            else:
                val = request.POST.get(item['clave'], '').strip()
                nuevos_datos[item['clave']] = val
            contexto[item['clave']] = val
            
        cliente.datos_extra.update(nuevos_datos)
        cliente.save(update_fields=['datos_extra'])
        
        doc.render(contexto)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        nombre = request.POST.get('nombre_archivo_salida', '').strip() or f"{plantilla.nombre} - {cliente.nombre_empresa}"
        if not nombre.lower().endswith('.docx'): nombre += ".docx"

        c_contratos, _ = Carpeta.objects.get_or_create(nombre="Contratos Generados", cliente=cliente, defaults={'es_expediente': False})
        nuevo = Documento(cliente=cliente, carpeta=c_contratos, nombre_archivo=nombre, subido_por=request.user)
        nuevo.archivo.save(nombre, ContentFile(buffer.getvalue()))
        nuevo.save()
        registrar_bitacora(request.user, cliente, 'generacion', f"Generó el contrato '{nombre}' desde la plantilla '{plantilla.nombre}'.")
        return redirect('visor_docx', documento_id=nuevo.id)

    return render(request, 'generador/llenar.html', {'cliente': cliente, 'plantilla': plantilla, 'variables': formulario})

@login_required
def visor_docx(request, documento_id):
    doc = get_object_or_404(Documento, id=documento_id)
    html = ""
    if doc.nombre_archivo.endswith('.docx'):
        try:
            html = mammoth.convert_to_html(io.BytesIO(doc.archivo.read())).value
        except Exception as e:
            logger.error(f"Error visualizando DOCX {documento_id}: {e}")
    return render(request, 'generador/visor.html', {'doc': doc, 'contenido_html': html})

@login_required
def subir_plantilla(request):
    if request.user.rol == 'admin' and request.method == 'POST':
        Plantilla.objects.create(nombre=request.POST.get('nombre'), archivo=request.FILES.get('archivo'))
    return redirect('dashboard')

@login_required
@requiere_permiso('access_disenador')
def diseñador_plantillas(request):
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        archivo = request.FILES.get('archivo_base')
        data_reemplazos = request.POST.get('reemplazos')

        if archivo and nombre:
            try:
                doc = DocumentoWord(archivo)
                
                if data_reemplazos:
                    lista = json.loads(data_reemplazos)
                    for item in lista:
                        original = item.get('texto_original', '')
                        variable_raw = item.get('variable', '').strip()
                        
                        if original and variable_raw:
                            variable_clean = variable_raw.lower().replace(' ', '_').replace('-', '_')
                            variable_clean = re.sub(r'[^a-z0-9_]', '', variable_clean)
                            variable_formateada = "{{ " + variable_clean + " }}"
                            reemplazar_preservando_estilo(doc, original, variable_formateada)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                nombre_archivo = nombre if nombre.endswith('.docx') else f"{nombre}.docx"
                
                nueva_plantilla = Plantilla(nombre=nombre)
                nueva_plantilla.archivo.save(nombre_archivo, ContentFile(buffer.getvalue()))
                nueva_plantilla.save()
                
                messages.success(request, f"¡Plantilla '{nombre}' guardada con estilos originales!")
                
            except Exception as e:
                logger.error(f"Error en diseñador: {e}")
                messages.error(request, f"Error procesando el archivo: {str(e)}")
            
            return redirect('dashboard')
            
    return render(request, 'generador/diseñador.html', {
        'glosario': VariableEstandar.objects.all().order_by('clave')
    })

@login_required
def previsualizar_word_raw(request):
    if request.method == 'POST' and request.FILES.get('archivo'):
        try:
            f = request.FILES['archivo']
            result = mammoth.convert_to_html(f)
            return JsonResponse({'html': result.value})
        except Exception as e:
            logger.error(f"Error en preview raw: {e}")
            return JsonResponse({'status': 'error', 'msg': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'msg': 'No se envió archivo'}, status=400)

@login_required
def crear_variable_api(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            clave_raw = data.get('clave', '')
            descripcion = data.get('descripcion', '')
            tipo = data.get('tipo', 'texto')

            if not clave_raw: 
                return JsonResponse({'status': 'error', 'msg': 'Falta la clave'}, status=400)

            clave_clean = clave_raw.lower().strip().replace(' ', '_').replace('-', '_')
            clave_clean = re.sub(r'[^a-z0-9_]', '', clave_clean)

            variable, created = VariableEstandar.objects.get_or_create(
                clave=clave_clean, 
                defaults={'descripcion': descripcion, 'tipo': tipo}
            )
            
            return JsonResponse({
                'status': 'ok', 
                'id': str(variable.id), 
                'clave': variable.clave,
                'created': created
            })
        except Exception as e:
            logger.error(f"Error creando variable API: {e}")
            return JsonResponse({'status': 'error', 'msg': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'msg': 'Método no permitido'}, status=405)

@login_required
def api_convertir_html(request):
    if request.method == 'POST':
        try:
            try: 
                data = json.loads(request.body)
                html_content = data.get('html', '')
            except: 
                html_content = request.POST.get('html', '')
                
            if not html_content: 
                return JsonResponse({'error': 'No content'}, status=400)
            
            TAGS_PDF_PERMITIDOS = [
                'p', 'br', 'strong', 'b', 'em', 'i', 'u', 's',
                'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                'table', 'thead', 'tbody', 'tfoot', 'tr', 'th', 'td',
                'div', 'span', 'section', 'article', 'header', 'footer',
                'ul', 'ol', 'li', 'dl', 'dt', 'dd',
                'img', 'a', 'hr', 'blockquote', 'pre', 'code',
            ]
            ATTRS_PDF_PERMITIDOS = {
                '*': ['style', 'class', 'id'],
                'a': ['href', 'target'],
                'img': ['src', 'alt', 'width', 'height'],
                'td': ['colspan', 'rowspan'],
                'th': ['colspan', 'rowspan'],
            }
            
            html_limpio = bleach.clean(
                html_content, 
                tags=TAGS_PDF_PERMITIDOS,
                attributes=ATTRS_PDF_PERMITIDOS,
                strip=True,
            )

            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="documento_diseñado.pdf"'
            base_url = request.build_absolute_uri('/')
            weasyprint.HTML(string=html_limpio, base_url=base_url).write_pdf(response)
            return response
        except Exception as e:
            logger.error(f"Error api_convertir_html: {e}")
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Only POST allowed'}, status=405)

# ==========================================
# 7. COTIZACIONES Y SERVICIOS
# ==========================================

@login_required
@requiere_permiso('access_cotizaciones')
def gestion_servicios(request):
    servicios = Servicio.objects.all().order_by('nombre')
    return render(request, 'cotizaciones/servicios.html', {'servicios': servicios})

@login_required
@requiere_permiso('access_cotizaciones')
def guardar_servicio(request):
    if request.method == 'POST':
        s_id = request.POST.get('servicio_id')
        s = get_object_or_404(Servicio, id=s_id) if s_id else Servicio()
        
        s.nombre = request.POST.get('nombre')
        s.descripcion = request.POST.get('descripcion')
        s.precio_base = request.POST.get('precio')
        
        nombres = request.POST.getlist('campo_nombre[]')
        valores = request.POST.getlist('campo_valor[]')
        
        estructura = []
        for nombre, valor in zip(nombres, valores):
            if nombre.strip():
                estructura.append({'nombre': nombre.strip(), 'valor': valor.strip()})
        
        s.campos_dinamicos = estructura
        s.save()
        messages.success(request, "Servicio actualizado correctamente.")
    return redirect('gestion_servicios')

@login_required
@requiere_permiso('access_cotizaciones')
def eliminar_servicio(request, servicio_id):
    get_object_or_404(Servicio, id=servicio_id).delete()
    return redirect('gestion_servicios')

@login_required
@requiere_permiso('access_cotizaciones')
def lista_cotizaciones(request):
    cotizaciones = Cotizacion.objects.select_related('creado_por', 'cliente_convertido').order_by('-fecha_creacion')
    return render(request, 'cotizaciones/lista.html', {'cotizaciones': cotizaciones})

@login_required
@requiere_permiso('access_cotizaciones')
@transaction.atomic
def nueva_cotizacion(request):
    if request.method == 'POST':
        titulo = request.POST.get('titulo')
        
        prospecto_empresa = request.POST.get('prospecto_empresa')
        prospecto_nombre = request.POST.get('prospecto_nombre')
        prospecto_email = request.POST.get('prospecto_email')
        prospecto_telefono = request.POST.get('prospecto_telefono')
        prospecto_direccion = request.POST.get('prospecto_direccion')
        prospecto_cargo = request.POST.get('prospecto_cargo')
        validez = request.POST.get('validez_hasta')
        
        if not titulo:
            cliente_ref = prospecto_empresa if prospecto_empresa else prospecto_nombre
            titulo = f"Cotización para {cliente_ref}"

        porcentaje_str = request.POST.get('porcentaje_descuento', '0')
        try: porcentaje_descuento = Decimal(porcentaje_str)
        except: porcentaje_descuento = Decimal('0.00')

        aplica_iva = request.POST.get('aplica_iva') == 'on'
        tasa_str = request.POST.get('porcentaje_iva_personalizado', '16')
        try: tasa_iva = Decimal(tasa_str)
        except: tasa_iva = Decimal('16.00')

        cotizacion = Cotizacion.objects.create(
            titulo=titulo,
            prospecto_empresa=prospecto_empresa,
            prospecto_nombre=prospecto_nombre,
            prospecto_email=prospecto_email,
            prospecto_telefono=prospecto_telefono,
            prospecto_direccion=prospecto_direccion,
            prospecto_cargo=prospecto_cargo,
            porcentaje_descuento=porcentaje_descuento,
            validez_hasta=validez if validez else None,
            condiciones_pago=request.POST.get('condiciones_pago', '50_50'),
            tiempo_entrega=request.POST.get('tiempo_entrega', '30_dias'),
            aplica_iva=aplica_iva,
            porcentaje_iva=tasa_iva,
            creado_por=request.user
        )

        servicios_ids = request.POST.getlist('servicios_seleccionados')
        cantidades = request.POST.getlist('cantidades')
        precios = request.POST.getlist('precios_personalizados')
        descripciones = request.POST.getlist('descripciones_personalizadas')

        items_to_create = []
        servicios_db = {str(s.id): s for s in Servicio.objects.filter(id__in=servicios_ids)}

        # Iteración segura por índices para evitar que un array más corto rompa el guardado
        for i in range(len(servicios_ids)):
            s_id = servicios_ids[i]
            
            if s_id and s_id in servicios_db:
                servicio = servicios_db[s_id]
                
                # Extracción segura con fallback a valores por defecto
                cantidad_str = cantidades[i] if i < len(cantidades) else '1'
                precio_str = precios[i] if i < len(precios) else '0'
                desc = descripciones[i] if i < len(descripciones) else ''
                
                try: 
                    cantidad = int(cantidad_str)
                except ValueError: 
                    cantidad = 1
                    
                try: 
                    precio_u = Decimal(precio_str)
                except ValueError: 
                    precio_u = Decimal('0.00')
                
                # Calcular subtotal manualmente para que bulk_create lo registre
                subtotal_calculado = Decimal(cantidad) * precio_u
                
                items_to_create.append(ItemCotizacion(
                    cotizacion=cotizacion,
                    servicio=servicio,
                    cantidad=cantidad,
                    precio_unitario=precio_u,
                    subtotal=subtotal_calculado,
                    descripcion_personalizada=desc
                ))
        
        # Guardar todo en bloque
        if items_to_create:
            ItemCotizacion.objects.bulk_create(items_to_create)
            
        # Recalcular totales generales de la cotización
        cotizacion.calcular_totales()

        messages.success(request, 'Cotización creada exitosamente.')
        return redirect('detalle_cotizacion', cotizacion_id=cotizacion.id)

    servicios = Servicio.objects.all()
    return render(request, 'cotizaciones/crear.html', {'servicios': servicios})

@login_required
def detalle_cotizacion(request, cotizacion_id):
    c = get_object_or_404(
        Cotizacion.objects.prefetch_related('items__servicio'), 
        id=cotizacion_id
    )
    
    recalculo_necesario = False
    if c.total == 0 and c.items.exists():
        recalculo_necesario = True
    
    for item in c.items.all():
        if item.subtotal == 0 and item.precio_unitario > 0:
            item.subtotal = item.cantidad * item.precio_unitario
            item.save()
            recalculo_necesario = True

    if recalculo_necesario:
        c.calcular_totales()
        c.refresh_from_db()

    # ---> INTELIGENCIA DE SUCURSALES (Idéntica a la lógica del Logo) <---
    sucursales_sugeridas = []
    if c.prospecto_empresa:
        palabras = c.prospecto_empresa.upper().split()
        
        if len(palabras) > 0:
            palabras_genericas = [
                'GRUPO', 'OPERADORA', 'COMERCIALIZADORA', 'EL', 'LA', 'LOS', 'LAS', 
                'CORPORATIVO', 'CONSORCIO', 'GASTRONOMIA', 'SERVICIOS', 'CONSTRUCTORA', 
                'PROMOTORA', 'PROVEEDORA', 'DISTRIBUIDORA', 'INMOBILIARIA', 'TRANSPORTES', 
                'LOGISTICA', 'RESTAURANTE', 'HOTEL', 'CLINICA', 'HOSPITAL', 'INSTITUTO', 
                'COLEGIO', 'AGENCIA', 'DESPACHO', 'ASOCIACION', 'SOCIEDAD', 'SISTEMAS', 
                'INDUSTRIAS', 'ADMINISTRADORA', 'CENTRO', 'FABRICA', 'PRODUCTORA'
            ]
            
            # Si empieza con palabra genérica, buscamos coincidencias con las primeras dos palabras
            if palabras[0] in palabras_genericas and len(palabras) > 1:
                clave_busqueda = f"{palabras[0]} {palabras[1]}"
            else:
                # Si es marca propia, buscamos coincidencias solo con la primera palabra
                clave_busqueda = palabras[0]
            
            if len(clave_busqueda) > 3:
                # Traemos TODAS las sucursales sin límite
                sucursales_sugeridas = Cliente.objects.filter(nombre_empresa__istartswith=clave_busqueda).order_by('nombre_empresa')

    return render(request, 'cotizaciones/detalle.html', {
        'c': c, 
        'plantillas_ws': PlantillaMensaje.objects.filter(tipo='whatsapp'),
        'sucursales_sugeridas': sucursales_sugeridas
    })

@login_required
def generar_pdf_cotizacion(request, cotizacion_id):
    c = get_object_or_404(Cotizacion.objects.prefetch_related('items__servicio'), id=cotizacion_id)
    
    if c.total == 0 and c.items.exists():
        c.calcular_totales() 
        c.refresh_from_db()  

    return generar_pdf_response(
        request, 
        'cotizaciones/pdf_template.html', 
        {'c': c}, 
        f"Cotizacion_{c.id}.pdf"
    )

@login_required
@transaction.atomic
def convertir_a_cliente(request, cotizacion_id):
    c = get_object_or_404(Cotizacion, id=cotizacion_id)

    if request.method != 'POST':
        return redirect('detalle_cotizacion', cotizacion_id=c.id)
    
    if c.cliente_convertido:
        messages.warning(request, f"Esta cotización ya es un cliente.")
        return redirect('detalle_cliente', cliente_id=c.cliente_convertido.id)

    # Limpiar items no seleccionados
    items_aceptados_ids = request.POST.getlist('items_seleccionados')
    items_a_borrar = ItemCotizacion.objects.filter(cotizacion=c).exclude(id__in=items_aceptados_ids)
    if items_a_borrar.exists():
        items_a_borrar.delete()
        c.calcular_totales()
        c.refresh_from_db()

    # ---> 1. OBTENER SUCURSALES SELECCIONADAS <---
    sucursales_ids = request.POST.getlist('sucursales_seleccionadas')
    clientes_afectados = list(Cliente.objects.filter(id__in=sucursales_ids))

    if not clientes_afectados:
        # Si no seleccionó ninguna (o es cliente nuevo), creamos uno normal
        nombre_busqueda = c.prospecto_empresa if c.prospecto_empresa else c.prospecto_nombre
        cli = Cliente.objects.filter(nombre_empresa__iexact=nombre_busqueda).first()
        if not cli:
            cli = Cliente.objects.create(
                nombre_empresa=nombre_busqueda,
                nombre_contacto=c.prospecto_nombre,
                email=c.prospecto_email,
                telefono=c.prospecto_telefono,
                datos_extra={'direccion': c.prospecto_direccion, 'cargo': c.prospecto_cargo}
            )
            if request.user.rol != 'admin':
                request.user.clientes_asignados.add(cli)
        clientes_afectados.append(cli)

    # El "Cliente Matriz" (al que se le cobra) será el primero de la lista
    cliente_principal = clientes_afectados[0]

    # ---> 2. FINANZAS: SE COBRA UNA SOLA VEZ (AL CLIENTE PRINCIPAL) <---
    monto_final = c.total_con_iva if c.aplica_iva else c.total
    hoy = timezone.now().date()
    
    dias_plazo = 15
    if c.tiempo_entrega == '30_dias': dias_plazo = 30
    elif c.tiempo_entrega == '60_dias': dias_plazo = 60
    elif c.tiempo_entrega == '90_dias': dias_plazo = 90
        
    fecha_final_proyecto = hoy + timedelta(days=dias_plazo)

    if c.condiciones_pago == '50_50':
        mitad = monto_final / Decimal(2)
        CuentaPorCobrar.objects.create(cliente=cliente_principal, cotizacion=c, concepto=f"50% Anticipo - {c.titulo}", monto_total=mitad, saldo_pendiente=mitad, fecha_vencimiento=hoy, estado='pendiente')
        CuentaPorCobrar.objects.create(cliente=cliente_principal, cotizacion=c, concepto=f"50% Liquidación - {c.titulo}", monto_total=mitad, saldo_pendiente=mitad, fecha_vencimiento=fecha_final_proyecto, estado='pendiente')
    elif c.condiciones_pago == '100_entrega':
        CuentaPorCobrar.objects.create(cliente=cliente_principal, cotizacion=c, concepto=f"Pago Contra Entrega - {c.titulo}", monto_total=monto_final, saldo_pendiente=monto_final, fecha_vencimiento=fecha_final_proyecto, estado='pendiente')
    else: 
        CuentaPorCobrar.objects.create(cliente=cliente_principal, cotizacion=c, concepto=f"Pago de Contado - {c.titulo}", monto_total=monto_final, saldo_pendiente=monto_final, fecha_vencimiento=hoy, estado='pendiente')


    # ---> 3. OPERACIÓN: SE CREAN CARPETAS EN TODAS LAS SUCURSALES <---
    carpetas_seleccionadas = request.POST.getlist('carpetas_seleccionadas')
    
    # Renderizamos el PDF solo una vez en memoria
    html_string = render_to_string('cotizaciones/pdf_template.html', {'c': c})
    html = weasyprint.HTML(string=html_string, base_url=request.build_absolute_uri())
    pdf_content = html.write_pdf()
    
    nombre_safe = slugify(c.titulo or f"v1_{c.id}").replace("-", "_")
    nombre_archivo = f"Cotizacion_{c.id}_{nombre_safe}_FINAL.pdf"

    for cli_sucursal in clientes_afectados:
        # Crea solo las carpetas seleccionadas (respeta las que ya existan)
        for nombre_carpeta in carpetas_seleccionadas:
            Carpeta.objects.get_or_create(nombre=nombre_carpeta, cliente=cli_sucursal, defaults={'es_expediente': False})
        
        # Guarda el PDF en su respectiva carpeta
        carpeta_cotizaciones, _ = Carpeta.objects.get_or_create(nombre="Cotizaciones", cliente=cli_sucursal, defaults={'es_expediente': False})
        if not Documento.objects.filter(carpeta=carpeta_cotizaciones, nombre_archivo=nombre_archivo).exists():
            nuevo_doc = Documento(cliente=cli_sucursal, carpeta=carpeta_cotizaciones, nombre_archivo=nombre_archivo, subido_por=request.user)
            nuevo_doc.archivo.save(nombre_archivo, ContentFile(pdf_content))
            nuevo_doc.save()

    # 4. Finalizar cotización
    c.estado = 'aceptada'
    c.cliente_convertido = cliente_principal
    c.save()

    registrar_bitacora(request.user, cliente_principal, 'creacion', f"Convirtió cotización corporativa para {len(clientes_afectados)} sucursales.")
    messages.success(request, f"¡Trato corporativo cerrado! Se armaron los expedientes en {len(clientes_afectados)} sucursal(es).")
    return redirect('detalle_cliente', cliente_id=cliente_principal.id)

@login_required
def enviar_cotizacion_email(request, cotizacion_id):
    cotizacion = get_object_or_404(Cotizacion.objects.prefetch_related('items__servicio'), id=cotizacion_id)
    
    if request.method == 'POST':
        asunto = request.POST.get('asunto')
        mensaje_usuario = request.POST.get('mensaje')
        firma_nombre = request.POST.get('firma_nombre', FIRMA_NOMBRE_DEFAULT)
        firma_cargo = request.POST.get('firma_cargo', FIRMA_CARGO_DEFAULT)
        usar_logo_default = request.POST.get('usar_logo_default') == 'on'
        
        html_string = render_to_string('cotizaciones/pdf_template.html', {'c': cotizacion})
        html = weasyprint.HTML(string=html_string, base_url=request.build_absolute_uri())
        pdf_file = html.write_pdf()

        html_content = f"""
        <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <div style="padding: 20px;">
                    <p style="white-space: pre-line;">{mensaje_usuario}</p>
                    <br><br>
                    <div style="border-top: 1px solid #ddd; padding-top: 20px; display: flex; align-items: center;">
                        {'<img src="cid:logo_firma" style="width: 50px; height: 50px; border-radius: 50%; margin-right: 15px;">' if usar_logo_default else ''}
                        <div>
                            <strong style="font-size: 14px; color: #2D1B4B; display: block;">{firma_nombre}</strong>
                            <span style="font-size: 12px; color: #666;">{firma_cargo}</span>
                        </div>
                    </div>
                </div>
            </body>
        </html>
        """
        text_content = strip_tags(html_content)

        email = EmailMultiAlternatives(
            subject=asunto,
            body=text_content,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[cotizacion.prospecto_email],
            reply_to=[EMAIL_REPLY_TO] 
        )
        email.attach_alternative(html_content, "text/html")

        filename = f"Cotizacion_{cotizacion.id}.pdf"
        email.attach(filename, pdf_file, 'application/pdf')

        if usar_logo_default:
            logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'logo.png') 
            if os.path.exists(logo_path):
                with open(logo_path, 'rb') as f:
                    logo_data = f.read()
                    logo = MIMEImage(logo_data)
                    logo.add_header('Content-ID', '<logo_firma>')
                    email.attach(logo)

        email.send()
        messages.success(request, f'Correo enviado exitosamente a {cotizacion.prospecto_email}')
        
    return redirect('detalle_cotizacion', cotizacion_id=cotizacion_id)

@login_required
@requiere_permiso('access_cotizaciones')
def eliminar_cotizacion(request, cotizacion_id):
    cotizacion = get_object_or_404(Cotizacion, id=cotizacion_id)
    cotizacion_id_ref = cotizacion.id 
    cotizacion.delete()
    messages.success(request, f"La cotización #{cotizacion_id_ref} fue eliminada exitosamente.")
    return redirect('lista_cotizaciones')

# ==========================================
# 8. FINANZAS (OPTIMIZADO)
# ==========================================

@login_required
@requiere_permiso('access_finanzas')
def panel_finanzas(request):
    clientes_con_actividad = (
        Cliente.objects
        .filter(cuentas__isnull=False)
        .distinct()
        .prefetch_related('cuentas')
    )

    lista_clientes = []
    total_global_pendiente = 0
    total_global_cobrado = 0

    for cli in clientes_con_actividad:
        cuentas = list(cli.cuentas.all())
        
        deuda = sum(c.saldo_pendiente for c in cuentas)
        pagado = sum(c.monto_pagado for c in cuentas)
        pendientes_count = sum(1 for c in cuentas if c.estado != 'pagado')
        
        lista_clientes.append({
            'obj': cli,
            'deuda': deuda,
            'pagado': pagado,
            'pendientes_count': pendientes_count
        })
        
        total_global_pendiente += deuda
        total_global_cobrado += pagado

    return render(request, 'finanzas/panel.html', {
        'clientes': lista_clientes,
        'total_por_cobrar': total_global_pendiente,
        'total_cobrado': total_global_cobrado
    })

@login_required
def registrar_pago(request):
    if request.method == 'POST':
        pago = Pago.objects.create(
            cuenta_id=request.POST.get('cuenta_id'), 
            monto=Decimal(request.POST.get('monto')),
            metodo=request.POST.get('metodo'), 
            referencia=request.POST.get('referencia'), 
            registrado_por=request.user
        )
        registrar_bitacora(
            request.user, pago.cuenta.cliente, 'pago',
            f"Registró pago de ${pago.monto:,.2f} vía {pago.metodo}. Ref: {pago.referencia or 'N/A'}."
        )
    return redirect('panel_finanzas')

@login_required
def recibo_pago_pdf(request, pago_id):
    p = get_object_or_404(Pago.objects.select_related('cuenta__cliente'), id=pago_id)
    return generar_pdf_response(request, 'finanzas/recibo_template.html', {'p': p}, f"Recibo_{p.id}.pdf")

@login_required
def eliminar_finanza(request, id):
    if request.user.rol != 'admin':
        messages.error(request, "Acceso denegado. Solo el Administrador puede eliminar registros financieros.")
        return redirect('panel_finanzas')
    
    cx = get_object_or_404(CuentaPorCobrar, id=id)
    registrar_bitacora(request.user, cx.cliente, 'eliminacion', f"Eliminó el registro financiero '{cx.concepto}' (${cx.monto_total:,.2f}).")
    cx.delete()
    messages.success(request, "Registro financiero eliminado correctamente.")
    return redirect('panel_finanzas')

@login_required
def finanzas_cliente(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    cuentas = cliente.cuentas.all().select_related('cotizacion').order_by('-fecha_vencimiento')
    
    proyectos = {}
    
    for cx in cuentas:
        clave = cx.cotizacion if cx.cotizacion else "Otros Cargos"
        
        if clave not in proyectos:
            proyectos[clave] = {
                'titulo': cx.cotizacion.titulo if cx.cotizacion else "Cargos Generales",
                'folio': cx.cotizacion.id if cx.cotizacion else None,
                'pagos': [],
                'total_proyecto': 0,
                'pendiente_proyecto': 0,
                'estado_general': 'completado' 
            }
        
        proyectos[clave]['pagos'].append(cx)
        proyectos[clave]['total_proyecto'] += cx.monto_total
        proyectos[clave]['pendiente_proyecto'] += cx.saldo_pendiente
        
        if cx.saldo_pendiente > 0:
            proyectos[clave]['estado_general'] = 'pendiente'

    return render(request, 'finanzas/detalle_cliente.html', {
        'cliente': cliente,
        'proyectos': proyectos
    })

@login_required
def generar_orden_cobro(request, cuenta_id, tipo_pago):
    cuenta = get_object_or_404(CuentaPorCobrar.objects.select_related('cotizacion', 'cliente'), id=cuenta_id)
    cotizacion = cuenta.cotizacion
    
    datos_bancarios = {
        'banco': request.GET.get('banco', 'BBVA'),
        'cuenta': request.GET.get('cuenta_num', ''),
        'clabe': request.GET.get('clabe', ''),
        'titular': request.GET.get('titular', '')
    }

    total_proyecto = cuenta.monto_total 
    
    if tipo_pago == 'anticipo':
        titulo_doc = "ORDEN DE PAGO - ANTICIPO"
        monto_a_pagar = total_proyecto / Decimal(2)
        nota = "Concepto: 50% de anticipo para inicio de gestiones administrativas."
        porcentaje_pago = 50
    else: 
        titulo_doc = "ORDEN DE PAGO - LIQUIDACIÓN"
        monto_a_pagar = cuenta.saldo_pendiente 
        nota = "Concepto: Pago final contra entrega de resultados."
        porcentaje_pago = 100 if cuenta.monto_pagado == 0 else 50 

    context = {
        'cuenta': cuenta,
        'c': cotizacion,
        'titulo_doc': titulo_doc,
        'monto_a_pagar': monto_a_pagar,
        'nota': nota,
        'tipo_pago': tipo_pago,
        'porcentaje_pago': porcentaje_pago,
        'banco': datos_bancarios,
        'fecha_emision': timezone.now()
    }

    return generar_pdf_response(request, 'finanzas/orden_cobro_pdf.html', context, f"Cobro_{tipo_pago}_{cuenta.cliente.nombre_empresa}.pdf")

# ==========================================
# 9. AGENDA
# ==========================================

@login_required
@requiere_permiso('access_agenda')
def agenda_legal(request):
    hoy = timezone.now()
    proximas = Evento.objects.filter(
        tipo='audiencia', inicio__gte=hoy, usuario=request.user
    ).select_related('cliente').order_by('inicio')[:5]
    
    clientes = Cliente.objects.all() if request.user.rol == 'admin' else request.user.clientes_asignados.all()
    return render(request, 'agenda/calendario.html', {'clientes': clientes, 'proximas_audiencias': proximas})

@login_required
def api_eventos(request):
    if not request.user.access_agenda: 
        return JsonResponse([], safe=False)
    start, end = request.GET.get('start'), request.GET.get('end')
    qs = Evento.objects.filter(inicio__range=[start, end]).select_related('cliente')
    
    if request.user.rol != 'admin': 
        qs = qs.filter(Q(usuario=request.user) | Q(cliente__in=request.user.clientes_asignados.all()))
    
    eventos = []
    for e in qs:
        titulo = f"{e.cliente.nombre_empresa}: {e.titulo}" if e.cliente else e.titulo
        eventos.append({'id': e.id, 'title': titulo, 'start': e.inicio.isoformat(), 'end': e.fin.isoformat() if e.fin else None, 'backgroundColor': e.color_hex, 'extendedProps': {'descripcion': e.descripcion, 'tipo': e.get_tipo_display()}})
    return JsonResponse(eventos, safe=False)

@login_required
def mover_evento_api(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            evento = get_object_or_404(Evento, id=data.get('id'))
            if request.user.rol != 'admin' and evento.usuario != request.user: 
                return JsonResponse({'status': 'error', 'msg': 'Sin permiso'})
            evento.inicio = data.get('start')
            if data.get('end'): evento.fin = data.get('end')
            evento.save()
            return JsonResponse({'status': 'ok'})
        except Exception as e: 
            logger.error(f"Error moviendo evento: {e}")
            return JsonResponse({'status': 'error', 'msg': str(e)})
    return JsonResponse({'status': 'error'})

@login_required
def crear_evento(request):
    if request.method == 'POST':
        inicio = timezone.make_aware(timezone.datetime.strptime(f"{request.POST.get('fecha')} {request.POST.get('hora')}", "%Y-%m-%d %H:%M"))
        cliente = get_object_or_404(Cliente, id=request.POST.get('cliente_id')) if request.POST.get('cliente_id') else None
        evento = Evento.objects.create(
            usuario=request.user, 
            titulo=request.POST.get('titulo'), 
            inicio=inicio, 
            tipo=request.POST.get('tipo'), 
            cliente=cliente, 
            descripcion=request.POST.get('descripcion')
        )
        if cliente:
            registrar_bitacora(request.user, cliente, 'agenda', f"Agendó el evento '{evento.titulo}' para el {inicio.strftime('%d/%m/%Y %H:%M')}.")
        messages.success(request, "Evento agendado.")
    return redirect('agenda_legal')

@login_required
def eliminar_evento(request, evento_id):
    evento = get_object_or_404(Evento, id=evento_id)
    if request.user.rol == 'admin' or evento.usuario == request.user:
        if evento.cliente:
            registrar_bitacora(request.user, evento.cliente, 'eliminacion', f"Eliminó el evento de agenda '{evento.titulo}'.")
        evento.delete()
        return JsonResponse({'status': 'ok'})
    return JsonResponse({'status': 'error'}, status=403)

@login_required
def eliminar_plantilla(request, plantilla_id):
    if request.user.rol != 'admin':
        messages.error(request, "No tienes permisos.")
        return redirect('dashboard')
        
    plantilla = get_object_or_404(Plantilla, id=plantilla_id)
    nombre = plantilla.nombre
    plantilla.archivo.delete() 
    plantilla.delete()
    
    messages.success(request, f"Plantilla '{nombre}' eliminada.")
    return redirect(request.META.get('HTTP_REFERER', 'dashboard'))

# ==========================================
# 10. GESTIÓN DE CARGA EXTERNA & UTILES
# ==========================================

@login_required
def subir_archivo_requisito(request, carpeta_id):
    if request.method == 'POST':
        carpeta_origen = get_object_or_404(Carpeta, id=carpeta_id)
        cliente = carpeta_origen.cliente
        archivo = request.FILES.get('archivo')
        nombre_requisito = request.POST.get('nombre_requisito')
        fecha_vencimiento = request.POST.get('fecha_vencimiento')

        if archivo and nombre_requisito:
            try:
                # ---> SE ELIMINÓ LA RESTRICCIÓN DE MAGIC PARA ADMITIR TODO TIPO DE ARCHIVOS <---
                
                anio_actual = timezone.now().year
                ext = archivo.name.split('.')[-1] if '.' in archivo.name else 'pdf'
                nuevo_nombre_formal = f"{nombre_requisito} {cliente.nombre_empresa} {anio_actual}.{ext}"

                carpetas_destino = []
                todas_carpetas = cliente.carpetas_drive.all()

                for carpeta in todas_carpetas:
                    requisitos_carpeta = carpeta.obtener_detalle_cumplimiento()
                    if requisitos_carpeta:
                        for req in requisitos_carpeta:
                            if req['nombre'] == nombre_requisito:
                                carpetas_destino.append(carpeta)
                                break
                
                if not carpetas_destino:
                    carpetas_destino.append(carpeta_origen)

                count = 0
                for carpeta_target in carpetas_destino:
                    Documento.objects.filter(
                        carpeta=carpeta_target, 
                        nombre_archivo__istartswith=nombre_requisito
                    ).delete()

                    nuevo_doc = Documento(
                        cliente=cliente,
                        carpeta=carpeta_target,
                        archivo=archivo,
                        nombre_archivo=nuevo_nombre_formal,
                        subido_por=request.user
                    )

                    if fecha_vencimiento: 
                        nuevo_doc.fecha_vencimiento = fecha_vencimiento
                        fecha_fin = datetime.strptime(fecha_vencimiento, '%Y-%m-%d').date()
                        alertas = [(20, '⚠️ Vence en 20 días'), (10, '🟠 Vence en 10 días'), (5, '🔴 URGENTE: Vence en 5 días')]
                        eventos_to_create = []
                        for dias_antes, prefijo in alertas:
                            fecha_alerta = fecha_fin - timedelta(days=dias_antes)
                            if fecha_alerta >= timezone.now().date():
                                eventos_to_create.append(Evento(
                                    cliente=cliente,
                                    usuario=request.user,
                                    titulo=f"{prefijo}: {nuevo_doc.nombre_archivo}",
                                    inicio=datetime.combine(fecha_alerta, datetime.min.time()),
                                    fin=datetime.combine(fecha_alerta, datetime.min.time()) + timedelta(hours=1),
                                    descripcion=f"Recordatorio automático de vencimiento para el documento: {nuevo_doc.nombre_archivo}"
                                ))
                        if eventos_to_create:
                            Evento.objects.bulk_create(eventos_to_create)

                    nuevo_doc.save()
                    count += 1

                registrar_bitacora(request.user, cliente, 'subida', f"Actualizó el requisito '{nombre_requisito}' en {count} carpeta(s) como '{nuevo_nombre_formal}'.")
                messages.success(request, f'✅ Archivo actualizado exitosamente en {count} carpeta(s) con el nombre: "{nuevo_nombre_formal}".')

            except Exception as e:
                logger.error(f"Error procesando archivo requisito: {e}")
                messages.error(request, f"Error al procesar el archivo: {e}")
        else:
            messages.error(request, 'Error: Faltan datos (archivo o nombre del requisito).')
            
        return redirect('detalle_cliente', cliente_id=cliente.id)
    
    return redirect('dashboard')

@login_required
def enviar_recordatorio_documentacion(request, cliente_id):
    cliente = get_object_or_404(Cliente.objects.prefetch_related('carpetas_drive'), id=cliente_id)
    
    faltantes_por_carpeta = {}
    total_faltantes = 0
    
    for carpeta in cliente.carpetas_drive.all():
        detalle = carpeta.obtener_detalle_cumplimiento()
        if detalle:
            items_rojos = [item['nombre'] for item in detalle if item['estado'] == 'missing']
            if items_rojos:
                faltantes_por_carpeta[carpeta.nombre] = items_rojos
                total_faltantes += len(items_rojos)
    
    if total_faltantes == 0:
        messages.success(request, "¡Este cliente ya tiene toda su documentación completa! No es necesario enviar recordatorios.")
        return redirect('detalle_cliente', cliente_id=cliente.id)

    asunto = f"Pendientes de Documentación - {cliente.nombre_empresa} - AppLegal"
    
    mensaje = f"""
Estimado(a) {cliente.nombre_contacto},

Esperamos que este correo le encuentre bien.

Le escribimos para darle seguimiento a su expediente de regularización. Para poder avanzar con los trámites ante las autoridades correspondientes, hemos detectado que aún tenemos algunos documentos pendientes de recibir.

A continuación, le compartimos el listado de los requisitos faltantes organizados por carpeta:
------------------------------------------------------------
"""

    for nombre_carpeta, documentos in faltantes_por_carpeta.items():
        mensaje += f"\n📂 {nombre_carpeta}:\n"
        for doc in documentos:
            mensaje += f"   [ ] {doc}\n"

    mensaje += f"""
------------------------------------------------------------

Le agradeceríamos mucho si pudiera compartirnos estos archivos a la brevedad posible, ya sea subiéndolos directamente a la plataforma o respondiendo a este correo.

Si tiene alguna duda sobre algún requisito en específico, quedamos totalmente a sus órdenes para apoyarle.

Atentamente,

Gestiones Cordpad
"""

    try:
        if cliente.email:
            send_mail(
                asunto,
                mensaje,
                settings.DEFAULT_FROM_EMAIL,
                [cliente.email],
                fail_silently=False,
            )
            messages.success(request, f"✅ Se envió el recordatorio a {cliente.email} con {total_faltantes} documentos faltantes.")
        else:
            messages.warning(request, "⚠️ El cliente no tiene un correo electrónico registrado.")
    except Exception as e:
        logger.error(f"Error enviando correo recordatorio: {e}")
        messages.error(request, f"❌ Error al enviar el correo: {str(e)}")

    return redirect('detalle_cliente', cliente_id=cliente.id)

# Asegúrate de tener 'Image' importado al principio de tu views.py, si no lo tienes agrégalo:
# from PIL import Image

@login_required
@requiere_permiso('access_qr')
def generador_qr(request):
    qr_url = None
    data = ""
    color_fill = "#2D1B4B"
    color_back = "#FFFFFF"

    if request.method == 'POST':
        data = request.POST.get('data')
        color_fill = request.POST.get('color_fill', '#2D1B4B')
        color_back = request.POST.get('color_back', '#FFFFFF')

        if data:
            # 1. Creamos el QR base (Usamos ERROR_CORRECT_H para que soporte tapar el centro)
            qr = qrcode.QRCode(
                version=1, 
                error_correction=qrcode.constants.ERROR_CORRECT_H, 
                box_size=10, 
                border=4
            )
            qr.add_data(data)
            qr.make(fit=True)
            
            # Generamos la imagen del QR y la convertimos a formato RGBA (para soportar transparencias)
            img_qr = qr.make_image(fill_color=color_fill, back_color=color_back).convert('RGBA')

            # 2. Cargamos el logo de Corpad
            logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'logo.png')
            
            if os.path.exists(logo_path):
                from PIL import Image # Importación segura
                
                logo = Image.open(logo_path)
                
                # Calculamos el tamaño del logo (máximo el 25% del QR para no romper la lectura)
                basewidth = int(img_qr.size[0] * 0.25)
                wpercent = (basewidth / float(logo.size[0]))
                hsize = int((float(logo.size[1]) * float(wpercent)))
                
                # Redimensionamos el logo
                logo = logo.resize((basewidth, hsize), Image.Resampling.LANCZOS)
                
                # Calculamos la posición central exacta
                pos_x = (img_qr.size[0] - logo.size[0]) // 2
                pos_y = (img_qr.size[1] - logo.size[1]) // 2
                
                # Pegamos el logo en el centro del QR usando la máscara alfa si es PNG
                if logo.mode in ('RGBA', 'LA') or (logo.mode == 'P' and 'transparency' in logo.info):
                    img_qr.paste(logo, (pos_x, pos_y), logo)
                else:
                    img_qr.paste(logo, (pos_x, pos_y))

            # 3. Convertimos la imagen final a Base64 para enviarla al template
            buffer = BytesIO()
            img_qr.save(buffer, format="PNG")
            img_str = base64.b64encode(buffer.getvalue()).decode()
            qr_url = f"data:image/png;base64,{img_str}"

    return render(request, 'generador_qr.html', {
        'qr_url': qr_url,
        'data_input': data,
        'color_fill': color_fill,
        'color_back': color_back
    })

@login_required
def buscar_cliente_api(request):
    query = request.GET.get('q', '')
    if len(query) < 2:
        return JsonResponse([], safe=False)
    
    clientes_encontrados = Cliente.objects.filter(
        Q(nombre_empresa__icontains=query) | 
        Q(nombre_contacto__icontains=query)
    )[:5]

    resultados = []
    for c in clientes_encontrados:
        direccion = ""
        cargo = ""
        if c.datos_extra and isinstance(c.datos_extra, dict):
            direccion = c.datos_extra.get('direccion', '')
            cargo = c.datos_extra.get('cargo', '')

        resultados.append({
            'prospecto_empresa': c.nombre_empresa,
            'prospecto_nombre': c.nombre_contacto,
            'prospecto_email': c.email,
            'prospecto_telefono': c.telefono,
            'prospecto_direccion': direccion,
            'prospecto_cargo': cargo
        })

    return JsonResponse(resultados, safe=False)

@login_required
def generar_link_externo(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    solicitud = SolicitudEnlace.objects.create(cliente=cliente)
    solicitud.fecha_expiracion = timezone.now() + timedelta(hours=72)
    solicitud.save()
    
    link = request.build_absolute_uri(f'/portal-cliente/{solicitud.id}/')
    registrar_bitacora(request.user, cliente, 'generacion', f"Generó link de carga externa (válido 72h).")
    messages.success(request, f"¡Link generado y expira en 72h! Copia y envía esto al cliente: {link}")
    return redirect('detalle_cliente', cliente_id=cliente.id)

def vista_publica_carga(request, token):
    solicitud = get_object_or_404(SolicitudEnlace, id=token, activa=True)
    
    if hasattr(solicitud, 'fecha_expiracion') and solicitud.fecha_expiracion and solicitud.fecha_expiracion < timezone.now():
        return HttpResponse('El link de carga ha expirado por seguridad. Por favor, solicite uno nuevo.', status=410)

    cliente = Cliente.objects.prefetch_related('carpetas_drive__documentos').get(id=solicitud.cliente.id)

    archivos_en_revision = set(ArchivoTemporal.objects.filter(solicitud=solicitud).values_list('nombre_requisito', flat=True))

    total_requisitos_unicos = set()
    requisitos_cumplidos_unicos = set()
    faltantes_reales = []
    faltantes_set = set()

    for carpeta in cliente.carpetas_drive.all():
        detalle = carpeta.obtener_detalle_cumplimiento()
        if detalle:
            for item in detalle:
                nombre = item['nombre']
                total_requisitos_unicos.add(nombre)
                
                if item['estado'] == 'ok':
                    requisitos_cumplidos_unicos.add(nombre)
                else:
                    if nombre not in archivos_en_revision:
                        if nombre not in faltantes_set:
                            faltantes_reales.append(nombre)
                            faltantes_set.add(nombre)
    
    tareas_completadas = len(requisitos_cumplidos_unicos | archivos_en_revision)
    total_tareas = len(total_requisitos_unicos)
    
    porcentaje = 0
    if total_tareas > 0:
        porcentaje = int((tareas_completadas / total_tareas) * 100)
    
    if not faltantes_reales and total_tareas > 0:
        porcentaje = 100

    documento_actual = faltantes_reales[0] if faltantes_reales else None

    if request.method == 'POST' and request.FILES.get('archivo'):
        requisito_a_subir = request.POST.get('requisito_objetivo')
        archivo_subido = request.FILES['archivo']
        
        # ---> SE ELIMINÓ LA RESTRICCIÓN DE MAGIC PARA ADMITIR TODO TIPO DE ARCHIVOS <---
        
        if requisito_a_subir:
            ArchivoTemporal.objects.create(
                solicitud=solicitud,
                archivo=archivo_subido,
                nombre_requisito=requisito_a_subir
            )
            return redirect('vista_publica_carga', token=token)
    
    return render(request, 'externo/portal_carga.html', {
        'cliente': cliente, 
        'documento_actual': documento_actual, 
        'faltantes_count': len(faltantes_reales),
        'porcentaje': porcentaje,
        'archivos_en_revision': archivos_en_revision
    })

@login_required
def aprobar_archivo_temporal(request, temp_id):
    temp = get_object_or_404(ArchivoTemporal, id=temp_id)
    cliente = temp.solicitud.cliente
    
    try:
        anio_actual = timezone.now().year
        ext = temp.archivo.name.split('.')[-1]
        nuevo_nombre_formal = f"{temp.nombre_requisito} {cliente.nombre_empresa} {anio_actual}.{ext}"
        
        carpetas_destino = []
        for carpeta in cliente.carpetas_drive.all():
            requisitos = carpeta.obtener_detalle_cumplimiento()
            if requisitos:
                for req in requisitos:
                    if req['nombre'] == temp.nombre_requisito:
                        carpetas_destino.append(carpeta)
                        break
        
        if not carpetas_destino:
            carpetas_destino.append(cliente.carpetas_drive.first())

        for carpeta_target in carpetas_destino:
            Documento.objects.filter(carpeta=carpeta_target, nombre_archivo__icontains=temp.nombre_requisito).delete()
            
            Documento.objects.create(
                cliente=cliente,
                carpeta=carpeta_target,
                archivo=temp.archivo,
                nombre_archivo=nuevo_nombre_formal,
                subido_por=request.user 
            )
            
        registrar_bitacora(request.user, cliente, 'aprobacion', f"Aprobó el archivo del portal externo: '{nuevo_nombre_formal}'.")
        temp.delete()
        messages.success(request, f"Aprobado y distribuido: {nuevo_nombre_formal}")
        
    except Exception as e:
        logger.error(f"Error aprobando archivo temporal: {e}")
        messages.error(request, f"Error al aprobar: {e}")

    return redirect('detalle_cliente', cliente_id=cliente.id)

@login_required
def rechazar_archivo_temporal(request, temp_id):
    temp = get_object_or_404(ArchivoTemporal, id=temp_id)
    nombre = temp.nombre_requisito
    cliente = temp.solicitud.cliente
    cliente_id = cliente.id
    
    registrar_bitacora(request.user, cliente, 'rechazo', f"Rechazó el archivo del portal externo: '{nombre}'.")
    temp.archivo.delete()
    temp.delete()
    
    messages.warning(request, f"❌ Documento rechazado y eliminado: {nombre}")
    return redirect('detalle_cliente', cliente_id=cliente_id)

# <--- FUNCIONES DE PREVIEW OPTIMIZADAS PARA S3 --->
@login_required
def preview_archivo(request, documento_id):
    doc = get_object_or_404(Documento, id=documento_id)
    ext = doc.nombre_archivo.split('.')[-1].lower()
    
    url_segura = doc.archivo.url
    
    data = {
        'nombre': doc.nombre_archivo,
        'url': url_segura,
        'tipo': 'desconocido',
        'html': ''
    }

    try:
        if ext in ['jpg', 'jpeg', 'png', 'gif', 'webp', 'svg']:
            data['tipo'] = 'imagen'
        elif ext in ['pdf', 'docx', 'xlsx', 'xls', 'csv', 'ppt', 'pptx']:
            url_codificada = urllib.parse.quote(url_segura)
            visor_url = f"https://docs.google.com/viewer?url={url_codificada}&embedded=true"
            data['tipo'] = 'pdf'
            data['url'] = visor_url
        elif ext in ['mp4', 'webm', 'ogg']:
            data['tipo'] = 'video'
        elif ext in ['mp3', 'wav']:
            data['tipo'] = 'audio'
        elif ext in ['txt', 'py', 'js', 'html', 'css', 'json', 'md']:
            data['tipo'] = 'texto'
            data['html'] = doc.archivo.read().decode('utf-8', errors='ignore') 
        else:
            data['tipo'] = 'descarga'

    except Exception as e:
        logger.error(f"Error generando preview {documento_id}: {e}")
        data['tipo'] = 'error'

    return JsonResponse(data)

@login_required
def obtener_preview_archivo(request, archivo_id):
    return preview_archivo(request, archivo_id)

# <--- DESCARGA DIRECTA OPTIMIZADA PARA S3 --->
@login_required
def descargar_archivo_oficial(request, archivo_id):
    doc = get_object_or_404(Documento, id=archivo_id)
    try:
        registrar_bitacora(request.user, doc.cliente, 'descarga', f"Descargó el archivo individual: '{doc.nombre_archivo}'.")
        
        response = HttpResponse(doc.archivo.read())
        
        nombre_codificado = urllib.parse.quote(doc.nombre_archivo)
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{nombre_codificado}"
        
        return response

    except Exception as e:
        logger.error(f"Error crítico en descarga: {str(e)}")
        messages.error(request, f"Error técnico al descargar: {str(e)}")
        return redirect('detalle_cliente', cliente_id=doc.cliente.id)

@login_required
def redactar_correo_autorizaciones(request, carpeta_id):
    carpeta = get_object_or_404(Carpeta, id=carpeta_id)
    cliente = carpeta.cliente
    
    acuse_id = request.GET.get('acuse_id')
    doc_acuse = None
    if acuse_id:
        doc_acuse = Documento.objects.filter(id=acuse_id).first()

    domain = URL_PORTAL
    logo_url = f"{domain}/static/img/logo.png"

    if request.method == 'POST':
        asunto = request.POST.get('asunto')
        mensaje_usuario = request.POST.get('mensaje')
        destinatario = request.POST.get('destinatario')
        firma_nombre = request.POST.get('firma_nombre', FIRMA_NOMBRE_DEFAULT)
        firma_cargo = request.POST.get('firma_cargo', FIRMA_CARGO_DEFAULT)
        
        acuse_id_post = request.POST.get('acuse_id_hidden')
        if acuse_id_post:
             doc_acuse = Documento.objects.filter(id=acuse_id_post).first()

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for doc in carpeta.documentos.all():
                if doc_acuse and doc.id == doc_acuse.id:
                    continue
                try:
                    zip_file.writestr(doc.nombre_archivo, doc.archivo.read())
                except Exception as e:
                    logger.warning(f"No se pudo adjuntar {doc.nombre_archivo}: {e}")
        buffer.seek(0)

        cuerpo_html = render_to_string('correo/email_body_universal.html', {
            'cliente': cliente,
            'tipo_correo': 'autorizaciones',
            'mensaje': mensaje_usuario,
            'lista_adjuntos': carpeta.documentos.exclude(id=doc_acuse.id if doc_acuse else None),
            'firma_nombre': firma_nombre,
            'firma_cargo': firma_cargo,
            'logo_url': logo_url,
        })
        
        text_body = strip_tags(cuerpo_html)
        
        email = EmailMultiAlternatives(
            subject=asunto,
            body=text_body,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[destinatario],
            cc=[request.user.email], 
            reply_to=[EMAIL_REPLY_TO] 
        )
        email.attach_alternative(cuerpo_html, "text/html")
        
        nombre_zip = f"Evidencias_{cliente.nombre_empresa}.zip"
        email.attach(nombre_zip, buffer.getvalue(), 'application/zip')
        
        if doc_acuse:
            try:
                email.attach(doc_acuse.nombre_archivo, doc_acuse.archivo.read(), 'application/pdf')
            except Exception as e:
                logger.warning(f"No se pudo adjuntar el físico del acuse: {e}")

        try:
            email.send()
            
            Bitacora.objects.create(
                usuario=request.user, 
                cliente=cliente, 
                accion='entrega_formal', 
                descripcion=f"Envió entrega final de: {carpeta.nombre}"
            )
            
            messages.success(request, f"✅ Correo enviado exitosamente a {destinatario} con Acuse y Evidencias.")
            return redirect('detalle_carpeta', cliente_id=cliente.id, carpeta_id=carpeta.id)
            
        except Exception as e:
            logger.error(f"Error enviando correo autorizaciones: {e}")
            messages.error(request, f"❌ Error al enviar: {e}")
    
    mensaje_plano = (
        f"Estimado(a) {cliente.nombre_contacto},\n\n"
        f"Por medio del presente le hago entrega formal de las autorizaciones originales gestionadas para {cliente.nombre_empresa}.\n\n"
        "Adjunto encontrará:\n"
        "1. Acuse de Recibo (PDF) con el detalle de vigencias.\n"
        "2. Archivo ZIP con las evidencias digitales de sus documentos.\n\n"
        "Quedo a sus órdenes."
    )
    
    return render(request, 'expedientes/redactar_correo.html', {
        'carpeta': carpeta,
        'cliente': cliente,
        'asunto': f"Entrega de Autorizaciones - {cliente.nombre_empresa}",
        'mensaje': mensaje_plano,
        'email_destino': cliente.email,
        'doc_acuse': doc_acuse 
    })

@login_required
def enviar_correo_universal(request, cliente_id, tipo_correo):
    cliente = get_object_or_404(Cliente.objects.prefetch_related('carpetas_drive'), id=cliente_id)
    
    domain = URL_PORTAL
    logo_url = f"{domain}/static/img/logo.png"

    context = {
        'cliente': cliente,
        'tipo_correo': tipo_correo,
        'destinatario': cliente.email or '',
        'firma_nombre': FIRMA_NOMBRE_DEFAULT,
        'firma_cargo': FIRMA_CARGO_DEFAULT,
        'logo_url': logo_url, 
    }

    if tipo_correo == 'cotizacion':
        cotizacion_id = request.GET.get('cotizacion_id') or request.POST.get('cotizacion_id')
        cotizacion = get_object_or_404(Cotizacion.objects.prefetch_related('items__servicio'), id=cotizacion_id)
        
        context.update({
            'cotizacion': cotizacion,
            'destinatario': cotizacion.prospecto_email or '',
            'asunto': f"Propuesta: {cotizacion.titulo}",
            'mensaje': (
                f"Estimado/a {cotizacion.prospecto_nombre},\n\n"
                f"Adjunto a este correo encontrará la propuesta detallada para el proyecto: {cotizacion.titulo}.\n\n"
                "Quedo a su entera disposición para cualquier duda."
            ),
            'url_cancelar': reverse('detalle_cotizacion', args=[cotizacion.id]),
        })

    elif tipo_correo == 'autorizaciones':
        carpeta_id = request.GET.get('carpeta_id') or request.POST.get('carpeta_id')
        carpeta = get_object_or_404(Carpeta.objects.prefetch_related('documentos'), id=carpeta_id)
        lista_adjuntos = carpeta.documentos.all()
        
        context.update({
            'carpeta': carpeta,
            'lista_adjuntos': lista_adjuntos,
            'total_adjuntos': lista_adjuntos.count(),
            'asunto': f"Entrega de Autorizaciones - {cliente.nombre_empresa}",
            'mensaje': (
                f"Estimado(a) {cliente.nombre_contacto},\n\n"
                f"Por medio del presente le hago entrega de las autorizaciones liberadas para {cliente.nombre_empresa}.\n\n"
                "Adjunto encontrará un archivo ZIP con todos los documentos digitales para su resguardo.\n\n"
                "Quedo a sus órdenes."
            ),
            'url_cancelar': reverse('detalle_carpeta', args=[cliente.id, carpeta.id]),
        })

    elif tipo_correo == 'recordatorio':
        faltantes_por_carpeta = {}
        total_faltantes = 0
        
        for carpeta in cliente.carpetas_drive.all():
            detalle = carpeta.obtener_detalle_cumplimiento()
            if detalle:
                items_rojos = [item['nombre'] for item in detalle if item['estado'] == 'missing']
                if items_rojos:
                    faltantes_por_carpeta[carpeta.nombre] = items_rojos
                    total_faltantes += len(items_rojos)
        
        if total_faltantes == 0:
            messages.success(request, "¡Este cliente ya tiene toda su documentación completa!")
            return redirect('detalle_cliente', cliente_id=cliente.id)

        solicitud = SolicitudEnlace.objects.create(cliente=cliente)
        solicitud.fecha_expiracion = timezone.now() + timedelta(hours=72)
        solicitud.save()
        link_carga = request.build_absolute_uri(f'/portal-cliente/{solicitud.id}/')

        context.update({
            'faltantes_por_carpeta': faltantes_por_carpeta,
            'total_faltantes': total_faltantes,
            'link_carga': link_carga,
            'asunto': f"Pendientes de Documentación - {cliente.nombre_empresa}",
            'mensaje': (
                f"Estimado(a) {cliente.nombre_contacto},\n\n"
                "Esperamos que se encuentre bien.\n\n"
                "Le escribimos para darle seguimiento a su expediente de regularización. "
                "Hemos detectado que aún tenemos algunos documentos pendientes de recibir.\n\n"
                "A continuación encontrará el listado detallado y un botón para subir los archivos directamente desde su dispositivo.\n\n"
                "Quedamos a sus órdenes para cualquier duda."
            ),
            'url_cancelar': reverse('detalle_cliente', args=[cliente.id]),
        })

    else:
        messages.error(request, "Tipo de correo no válido.")
        return redirect('detalle_cliente', cliente_id=cliente.id)

    if request.method == 'POST':
        destinatario = request.POST.get('destinatario')
        asunto = request.POST.get('asunto')
        mensaje = request.POST.get('mensaje')
        firma_nombre = request.POST.get('firma_nombre')
        firma_cargo = request.POST.get('firma_cargo')
        
        email_context = {
            'cliente': cliente,
            'tipo_correo': tipo_correo,
            'mensaje': mensaje,
            'firma_nombre': firma_nombre,
            'firma_cargo': firma_cargo,
            'logo_url': logo_url, 
        }

        try:
            pdf_bytes = None
            zip_buffer = None
            filename_adjunto = None

            if tipo_correo == 'cotizacion':
                cotizacion_id = request.POST.get('cotizacion_id')
                cotizacion = get_object_or_404(Cotizacion.objects.prefetch_related('items__servicio'), id=cotizacion_id)
                email_context['cotizacion'] = cotizacion
                
                html_string = render_to_string('cotizaciones/pdf_template.html', {'c': cotizacion})
                pdf_bytes = weasyprint.HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf()
                filename_adjunto = f"Cotizacion_{cotizacion.id}.pdf"

            elif tipo_correo == 'autorizaciones':
                carpeta_id = request.POST.get('carpeta_id')
                carpeta = get_object_or_404(Carpeta.objects.prefetch_related('documentos'), id=carpeta_id)
                lista_adjuntos = carpeta.documentos.all()
                email_context['lista_adjuntos'] = lista_adjuntos
                
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                    for doc in lista_adjuntos:
                        try:
                            zip_file.writestr(doc.nombre_archivo, doc.archivo.read())
                        except Exception as e:
                            logger.warning(f"Error comprimiendo {doc.nombre_archivo}: {e}")
                zip_buffer.seek(0)
                filename_adjunto = f"Autorizaciones_{cliente.nombre_empresa}_{timezone.now().date()}.zip"

            elif tipo_correo == 'recordatorio':
                solicitud = SolicitudEnlace.objects.filter(cliente=cliente, activa=True).last()
                if not solicitud:
                    solicitud = SolicitudEnlace.objects.create(cliente=cliente)
                    solicitud.fecha_expiracion = timezone.now() + timedelta(hours=72)
                    solicitud.save()
                link_carga = request.build_absolute_uri(f'/portal-cliente/{solicitud.id}/')
                
                email_context['faltantes_por_carpeta'] = context.get('faltantes_por_carpeta', {})
                email_context['link_carga'] = link_carga

            html_body = render_to_string('correo/email_body_universal.html', email_context)
            text_body = strip_tags(html_body) 

            email = EmailMultiAlternatives(
                subject=asunto,
                body=text_body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=[destinatario],
                reply_to=[EMAIL_REPLY_TO] 
            )
            
            if tipo_correo == 'autorizaciones':
                email.cc = [request.user.email]

            email.attach_alternative(html_body, "text/html")

            if pdf_bytes:
                email.attach(filename_adjunto, pdf_bytes, 'application/pdf')
            
            if zip_buffer:
                email.attach(filename_adjunto, zip_buffer.getvalue(), 'application/zip')

            email.send()
            messages.success(request, f"✅ Correo enviado exitosamente a {destinatario}")
            
            Bitacora.objects.create(
                usuario=request.user, cliente=cliente, 
                accion='envio_correo', 
                descripcion=f"Envió correo ({tipo_correo}) a {destinatario}: {asunto}"
            )

        except Exception as e:
            logger.error(f"Error enviando correo universal ({tipo_correo}): {e}")
            messages.error(request, f"❌ Error al enviar el correo: {str(e)}")

        if tipo_correo == 'cotizacion':
            return redirect('detalle_cotizacion', cotizacion_id=request.POST.get('cotizacion_id'))
        elif tipo_correo == 'autorizaciones':
            return redirect('detalle_carpeta', cliente_id=cliente.id, carpeta_id=request.POST.get('carpeta_id'))
        else:
            return redirect('detalle_cliente', cliente_id=cliente.id)

    return render(request, 'correo/enviar_correo_universal.html', context)

def reemplazar_preservando_estilo(doc, texto_original, texto_nuevo):
    if not texto_original or not texto_nuevo:
        return

    for p in doc.paragraphs:
        if texto_original in p.text:
            for run in p.runs:
                if texto_original in run.text:
                    run.text = run.text.replace(texto_original, texto_nuevo)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if texto_original in p.text:
                        for run in p.runs:
                            if texto_original in run.text:
                                run.text = run.text.replace(texto_original, texto_nuevo)

@login_required
def generar_contrato_final(request):
    if request.method == 'POST':
        try:
            plantilla_id = request.POST.get('plantilla_id')
            plantilla = get_object_or_404(Plantilla, id=plantilla_id)

            campos_ignorados = ['csrfmiddlewaretoken', 'plantilla_id', 'nombre_archivo_salida']
            contexto = {}

            for key, value in request.POST.items():
                if key not in campos_ignorados:
                    contexto[key] = value

            doc = DocxTemplate(io.BytesIO(plantilla.archivo.read()))
            doc.render(contexto)

            nombre_salida = request.POST.get('nombre_archivo_salida', 'Documento_Generado')
            nombre_salida = nombre_salida.strip().replace('/', '_').replace('\\', '_')
            if not nombre_salida.lower().endswith('.docx'):
                nombre_salida += '.docx'

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{nombre_salida}"'

            doc.save(response)
            return response

        except Exception as e:
            print(f"Error generando contrato: {str(e)}")
            return redirect('dashboard') 

    return redirect('dashboard')

@login_required
def preparar_entrega_autorizaciones(request, cliente_id, carpeta_id):
    carpeta = get_object_or_404(Carpeta.objects.prefetch_related('documentos'), id=carpeta_id)
    cliente = carpeta.cliente
    documentos = carpeta.documentos.all()

    observaciones_default = (
        "1. Las Licencias y Autorizaciones originales deberán ser exhibidas en un lugar visible dentro del establecimiento.\n"
        "2. Se recomienda realizar un respaldo digital adicional de estos archivos.\n"
        "3. Informar cualquier cambio o modificación a las autoridades pertinentes."
    )

    if request.method == 'POST':
        try:
            atencion = request.POST.get('atencion')
            cargo = request.POST.get('cargo') 
            observaciones = request.POST.get('observaciones')
            municipio = request.POST.get('municipio', 'Cuautitlán')
            estado = request.POST.get('estado', 'Estado de México')
            
            docs_procesados = []
            for doc in documentos:
                raw_detalle = request.POST.get(f'detalle_{doc.id}', '').strip()
                mostrar_vence = request.POST.get(f'vence_{doc.id}') == 'on'
                
                nombre_limpio = os.path.splitext(doc.nombre_archivo)[0]
                detalle_lista = [linea.strip() for linea in raw_detalle.split('\n') if linea.strip()]

                docs_procesados.append({
                    'nombre': nombre_limpio,
                    'detalle_lista': detalle_lista,
                    'mostrar_vence': mostrar_vence,
                    'fecha_vencimiento': doc.fecha_vencimiento
                })

            logo_b64 = ""
            firma_b64 = ""
            logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'logo.png')
            firma_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'firma_full.png')
            
            if os.path.exists(logo_path):
                with open(logo_path, "rb") as img_file:
                    logo_b64 = base64.b64encode(img_file.read()).decode('utf-8')
                    
            if os.path.exists(firma_path):
                with open(firma_path, "rb") as img_file:
                    firma_b64 = base64.b64encode(img_file.read()).decode('utf-8')

            context_pdf = {
                'cliente': cliente,
                'carpeta': carpeta,
                'documentos': docs_procesados,
                'atencion': atencion,
                'cargo': cargo, 
                'observaciones': observaciones,
                'municipio': municipio,
                'estado': estado,
                'fecha_emision': timezone.now(),
                'logo_b64': logo_b64,
                'firma_b64': firma_b64,
            }
            
            html_string = render_to_string('expedientes/acuse_pdf.html', context_pdf)
            html = weasyprint.HTML(string=html_string)
            pdf_content = html.write_pdf()

            nombre_acuse = f"ACUSE ENTREGA {cliente.nombre_empresa}.pdf"
            Documento.objects.filter(carpeta=carpeta, nombre_archivo__startswith="ACUSE ENTREGA").delete()
            
            nuevo_acuse = Documento(
                cliente=cliente,
                carpeta=carpeta,
                nombre_archivo=nombre_acuse,
                subido_por=request.user
            )
            nuevo_acuse.archivo.save(nombre_acuse, ContentFile(pdf_content))
            nuevo_acuse.save()

            registrar_bitacora(request.user, cliente, 'generacion', f"Generó el acuse de entrega para la carpeta '{carpeta.nombre}'.")

            url_redactar = reverse('redactar_correo_autorizaciones', kwargs={'carpeta_id': carpeta.id})
            return redirect(f"{url_redactar}?acuse_id={nuevo_acuse.id}")

        except Exception as e:
            logger.error(f"Error generando acuse: {e}")
            messages.error(request, f"Error al generar el acuse: {e}")

    return render(request, 'expedientes/configurar_entrega.html', {
        'carpeta': carpeta,
        'cliente': cliente,
        'documentos': documentos,
        'observaciones_default': observaciones_default
    })

@login_required
@requiere_permiso('access_gastos')
def modulo_gastos(request):
    if request.method == 'POST' and request.FILES.getlist('xml_files'):
        archivos = request.FILES.getlist('xml_files')
        guardados, errores, existentes = 0, 0, 0
        
        for f in archivos:
            if not f.name.lower().endswith('.xml'): continue
            try:
                datos = procesar_xml_factura(f)
                if FacturaGasto.objects.filter(uuid=datos['uuid']).exists():
                    existentes += 1
                    continue
                
                FacturaGasto.objects.create(
                    uuid=datos['uuid'],
                    fecha_emision=datos['fecha_emision'],
                    rfc_emisor=datos['rfc_emisor'],
                    nombre_emisor=datos['nombre_emisor'],
                    rfc_receptor=datos['rfc_receptor'],
                    nombre_receptor=datos['nombre_receptor'],
                    subtotal=datos['subtotal'],
                    total_impuestos=datos['total_impuestos'],
                    total=datos['total'],
                    moneda=datos['moneda'],
                    archivo_xml=f,
                    cargado_por=request.user
                )
                guardados += 1
            except Exception as e:
                logger.error(f"Error XML: {e}")
                errores += 1
        
        msj = f"Carga: {guardados} nuevos."
        if existentes: msj += f" {existentes} repetidos."
        if errores: msj += f" {errores} fallidos."
        messages.success(request, msj) if guardados else messages.warning(request, msj)
        return redirect('modulo_gastos')

    anio = request.GET.get('anio', timezone.now().year)
    mes = request.GET.get('mes', '')

    gastos = FacturaGasto.objects.filter(fecha_emision__year=anio)

    if mes and mes != '':
        gastos = gastos.filter(fecha_emision__month=mes)

    totales_db = gastos.aggregate(
        suma_total=Sum('total'), 
        suma_iva=Sum('total_impuestos'),
        suma_subtotal=Sum('subtotal') 
    )
    
    total_filtrado = totales_db['suma_total'] or 0
    iva_filtrado = totales_db['suma_iva'] or 0
    subtotal_filtrado = totales_db['suma_subtotal'] or 0 

    resumen_mes = gastos.annotate(mes_trunc=TruncMonth('fecha_emision')).values('mes_trunc').annotate(
        total_mes=Sum('total'),
        iva_mes=Sum('total_impuestos'),
        count=Count('id')
    ).order_by('mes_trunc')

    return render(request, 'finanzas/gastos.html', {
        'resumen_mes': resumen_mes,
        'total_anual': total_filtrado,
        'iva_anual': iva_filtrado,
        'subtotal_anual': subtotal_filtrado, 
        'anio_actual': int(anio),
        'mes_actual': mes,
        'ultimos_gastos': gastos.order_by('-fecha_emision')
    })
    
@login_required
def exportar_gastos_excel(request):
    if request.user.rol != 'admin': return redirect('dashboard')
    
    anio = request.GET.get('anio', timezone.now().year)
    mes = request.GET.get('mes', '')

    qs = FacturaGasto.objects.filter(fecha_emision__year=anio)
    nombre_archivo = f"Gastos_{anio}"

    if mes:
        qs = qs.filter(fecha_emision__month=mes)
        nombres_meses = ['', 'Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
        nombre_mes = nombres_meses[int(mes)]
        nombre_archivo += f"_{nombre_mes}"

    if not qs.exists():
        messages.warning(request, "No hay datos para exportar con esos filtros.")
        return redirect('modulo_gastos')

    data = list(qs.values(
        'fecha_emision', 'uuid', 'nombre_emisor', 'rfc_emisor', 
        'subtotal', 'total_impuestos', 'total', 'moneda'
    ))
    
    df = pd.DataFrame(data)
    
    df['fecha_emision'] = df['fecha_emision'].astype(str).str[:10]
    df.rename(columns={
        'fecha_emision': 'Fecha', 'uuid': 'UUID', 'nombre_emisor': 'Proveedor',
        'rfc_emisor': 'RFC', 'total_impuestos': 'IVA Trasladado', 'total': 'Total'
    }, inplace=True)

    response = HttpResponse(content_type='application/vnd.ms-excel')
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}.xlsx"'
    
    with pd.ExcelWriter(response, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Gastos')
        
    return response
    
@login_required
def crear_carpetas_especiales(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    
    if request.method == 'POST':
        carpetas_seleccionadas = request.POST.getlist('carpetas')
        
        for nombre in carpetas_seleccionadas:
            carpeta, creada = Carpeta.objects.get_or_create(
                nombre=nombre,
                cliente=cliente,
                defaults={'es_expediente': False}
            )
            if creada:
                registrar_bitacora(request.user, cliente, 'creacion', f"Se generó la carpeta del sistema: '{nombre}'.")
        
        messages.success(request, f"Se generaron o actualizaron {len(carpetas_seleccionadas)} carpetas especiales.")
        
    return redirect('detalle_cliente', cliente_id=cliente.id)

@login_required
def eliminar_carpetas_especiales(request, cliente_id):
    if not (request.user.can_delete_client or request.user.rol == 'admin'):
        messages.error(request, "No tienes permiso para eliminar carpetas.")
        return redirect('detalle_cliente', cliente_id=cliente_id)

    if request.method == 'POST':
        cliente = get_object_or_404(Cliente, id=cliente_id)
        alcance = request.POST.get('alcance', 'solo_este') # <-- 'solo_este' o 'todas_sucursales'
        
        nombres_semaforo = [
            'Autorizaciones liberadas', 'CARPETA ADMINISTRATIVA',
            'LICENCIA DE FUNCIONAMIENTO', 'PROGRAMA ESPECIFICO DE PROTECCIÓN CIVIL',
            'PROTECCIÓN CIVIL MUNICIPAL', 'PROTECCIÓN CIVIL ESTATAL',
            'MEDIO AMBIENTE', 'REGISTRO AMBIENTAL ESTATAL',
            'CEDULA DE ZONIFICACIÓN', 'LICENCIA DE USO DE SUELO'
        ]
        
        clientes_afectados = [cliente]
        
        # Si eligió MASIVO, buscamos todas las sucursales
        if alcance == 'todas_sucursales':
            palabras = cliente.nombre_empresa.upper().split()
            if len(palabras) > 0:
                palabras_genericas = [
                    'GRUPO', 'OPERADORA', 'COMERCIALIZADORA', 'EL', 'LA', 'LOS', 'LAS', 
                    'CORPORATIVO', 'CONSORCIO', 'GASTRONOMIA', 'SERVICIOS', 'CONSTRUCTORA', 
                    'PROMOTORA', 'PROVEEDORA', 'DISTRIBUIDORA', 'INMOBILIARIA', 'TRANSPORTES', 
                    'LOGISTICA', 'RESTAURANTE', 'HOTEL', 'CLINICA', 'HOSPITAL', 'INSTITUTO', 
                    'COLEGIO', 'AGENCIA', 'DESPACHO', 'ASOCIACION', 'SOCIEDAD', 'SISTEMAS', 
                    'INDUSTRIAS', 'ADMINISTRADORA', 'CENTRO', 'FABRICA', 'PRODUCTORA'
                ]
                if palabras[0] in palabras_genericas and len(palabras) > 1:
                    clave_busqueda = f"{palabras[0]} {palabras[1]}"
                else:
                    clave_busqueda = palabras[0]
                
                if len(clave_busqueda) > 3:
                    clientes_afectados = list(Cliente.objects.filter(nombre_empresa__istartswith=clave_busqueda))
        
        total_borradas = 0
        for cli in clientes_afectados:
            carpetas = Carpeta.objects.filter(cliente=cli, nombre__in=nombres_semaforo)
            count = carpetas.count()
            if count > 0:
                for c in carpetas:
                    c.delete()
                total_borradas += count
        
        if total_borradas > 0:
            registrar_bitacora(request.user, cliente, 'eliminacion', f"Eliminó {total_borradas} carpetas especiales en {len(clientes_afectados)} sucursal(es).")
            messages.success(request, f"¡Éxito! Se eliminaron {total_borradas} carpetas especiales en {len(clientes_afectados)} sucursal(es).")
        else:
            messages.warning(request, "No se encontraron carpetas especiales para eliminar.")
            
    return redirect('detalle_cliente', cliente_id=cliente.id)