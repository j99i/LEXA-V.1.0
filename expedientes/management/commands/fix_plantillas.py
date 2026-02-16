from django.core.management.base import BaseCommand
from docx import Document
from expedientes.models import Plantilla
import re

class Command(BaseCommand):
    help = 'Corrige variables con espacios en plantillas docx'

    def handle(self, *args, **kwargs):
        self.stdout.write("🔍 Escaneando plantillas...")
        
        for plantilla in Plantilla.objects.all():
            try:
                doc = Document(plantilla.archivo.path)
                corregido = False
                
                # Buscar en párrafos
                for p in doc.paragraphs:
                    for run in p.runs:
                        if '{{' in run.text and '}}' in run.text:
                            matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', run.text)
                            for match in matches:
                                if ' ' in match or '-' in match or match != match.lower():
                                    nueva = match.strip().replace(' ', '_').replace('-', '_').lower()
                                    nueva = re.sub(r'[^a-z0-9_]', '', nueva)
                                    run.text = run.text.replace(match, nueva)
                                    corregido = True
                                    self.stdout.write(f"  ✅ '{match}' → '{nueva}'")
                
                # Buscar en tablas también
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    if '{{' in run.text and '}}' in run.text:
                                        matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', run.text)
                                        for match in matches:
                                            if ' ' in match or '-' in match or match != match.lower():
                                                nueva = match.strip().replace(' ', '_').replace('-', '_').lower()
                                                nueva = re.sub(r'[^a-z0-9_]', '', nueva)
                                                run.text = run.text.replace(match, nueva)
                                                corregido = True
                                                self.stdout.write(f"  ✅ Tabla: '{match}' → '{nueva}'")

                if corregido:
                    doc.save(plantilla.archivo.path)
                    self.stdout.write(self.style.SUCCESS(f'💾 Plantilla "{plantilla.nombre}" guardada'))
                else:
                    self.stdout.write(f'✔️  "{plantilla.nombre}" - OK')
                    
            except Exception as e:
                self.stdout.write(self.style.ERROR(f'❌ Error en "{plantilla.nombre}": {e}'))
        
        self.stdout.write(self.style.SUCCESS("🏁 Proceso completado"))